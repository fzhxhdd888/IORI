# !/usr/bin/env python
# _*_ coding: utf-8 _*_
# @Author    : Feng Zhaohui
# @Time      : 2019/2/22
# @File      : report_handle.py
# @Funcyusa  :
# @Version   : 1.0

from openpyxl import Workbook, load_workbook
import os
import time
import global_element
from docx import Document
from docx.shared import Inches
import shutil
from win32com.client import Dispatch


# Judgement判断
def judgement_handle(result, lowlimit, highlimit):
    if result != 'NULL':
        if lowlimit == 'None' and highlimit == 'None':
            Judgement = 'Passed'
        elif lowlimit == 'None' and highlimit != 'None':
            if float(highlimit) - float(result) > 0:
                Judgement = 'Passed'
            else:
                Judgement = 'Failed'
        elif lowlimit != 'None' and highlimit == 'None':
            if float(result) - float(lowlimit) > 0:
                Judgement = 'Passed'
            else:
                Judgement = 'Failed'
        else:
            if float(result) - float(lowlimit) > 0 and float(highlimit) - float(result) > 0:
                Judgement = 'Passed'
            else:
                Judgement = 'Failed'
    else:
        Judgement = 'Inconclusive'
    return Judgement


# wlan fre calc
def fre_wlan_calc(channel):
    if int(channel) in range(1, 14):
        fre = int(channel) * 5 + 2407
    elif int(channel) == 14:
        fre = 2484
    elif int(channel) in range(34, 197):
        fre = 5000 + int(channel) * 5
    else:
        fre = ''

    return str(fre)


# wlan fre calc
def fre_bt2_calc(channel):
    try:
        if int(channel) in range(0, 79):
            fre = int(channel) * 1 + 2402
        else:
            fre = ''
    except:
        fre = ''

    return str(fre)


# GSM的DL_Freq计算
def freq_gsm_calc(channel):
    if global_element.Test_band == 'GSM850':
        if channel in range(128, 252):
            dlfreq = 869.2 + (channel - 128) * 0.2
            ulfreq = 824.2 + (channel - 128) * 0.2
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'GSM900':
        if channel in range(975, 1024):
            dlfreq = 925.2 + (channel - 975) * 0.2
            ulfreq = 880.2 + (channel - 975) * 0.2
        elif channel in range(0, 125):
            dlfreq = 935.0 + (channel - 0) * 0.2
            ulfreq = 890.0 + (channel - 0) * 0.2
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'DCS1800':
        if channel in range(512, 886):
            dlfreq = 1805.2 + (channel - 512) * 0.2
            ulfreq = 1710.2 + (channel - 512) * 0.2
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'PCS1900':
        if channel in range(512, 811):
            dlfreq = 1930.2 + (channel - 512) * 0.2
            ulfreq = 1850.2 + (channel - 512) * 0.2
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'Band II':
        if channel in range(9262, 9539):
            dlfreq = 1932.4 + (channel - 9262) * 0.2
            ulfreq = 1852.4 + (channel - 9262) * 0.2
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'Band IV':
        if channel in range(1312, 1514):
            dlfreq = 2112.4 + (channel - 1312) * 0.2
            ulfreq = 1712.4 + (channel - 1312) * 0.2
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'Band V':
        if channel in range(4132, 4234):
            dlfreq = 871.4 + (channel - 4132) * 0.2
            ulfreq = 826.4 + (channel - 4132) * 0.2
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 1':
        if channel in range(18025, 18576):
            dlfreq = 2112.5 + (channel - 18025) * 0.1
            ulfreq = 1922.5 + (channel - 18025) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band in ['FDD-LTE 2', 'CA_2C']:
        if channel in range(18607, 19194):
            dlfreq = 1930.7 + (channel - 18607) * 0.1
            ulfreq = 1850.7 + (channel - 18607) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 3':
        if channel in range(19207, 19944):
            dlfreq = 1805.7 + (channel - 19207) * 0.1
            ulfreq = 1710.7 + (channel - 19207) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 4':
        if channel in range(19957, 20394):
            dlfreq = 2110.7 + (channel - 19957) * 0.1
            ulfreq = 1710.7 + (channel - 19957) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band in ['FDD-LTE 5', 'CA_5B']:
        if channel in range(20407, 20644):
            dlfreq = 869.7 + (channel - 20407) * 0.1
            ulfreq = 824.7 + (channel - 20407) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 6':
        if channel in range(20675, 20726):
            dlfreq = 877.5 + (channel - 20675) * 0.1
            ulfreq = 832.5 + (channel - 20675) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band in ['FDD-LTE 7', 'CA_7C']:
        if channel in range(20775, 21426):
            dlfreq = 2622.5 + (channel - 20775) * 0.1
            ulfreq = 2502.5 + (channel - 20775) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 8':
        if channel in range(21457, 21794):
            dlfreq = 925.7 + (channel - 21457) * 0.1
            ulfreq = 880.7 + (channel - 21457) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 9':
        if channel in range(21825, 22126):
            dlfreq = 1847.4 + (channel - 21825) * 0.1
            ulfreq = 1752.4 + (channel - 21825) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 10':
        if channel in range(22175, 22726):
            dlfreq = 2112.5 + (channel - 22175) * 0.1
            ulfreq = 1712.5 + (channel - 22175) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 11':
        if channel in range(22775, 22926):
            dlfreq = 1478.4 + (channel - 22775) * 0.1
            ulfreq = 1430.4 + (channel - 22775) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band in ['FDD-LTE 12', 'CA_12B']:
        if channel in range(23017, 23174):
            dlfreq = 729.7 + (channel - 23017) * 0.1
            ulfreq = 699.7 + (channel - 23017) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 13':
        if channel in range(23205, 23256):
            dlfreq = 748.5 + (channel - 23205) * 0.1
            ulfreq = 779.5 + (channel - 23205) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 14':
        if channel in range(23305, 23356):
            dlfreq = 760.5 + (channel - 23305) * 0.1
            ulfreq = 790.5 + (channel - 23305) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 17':
        if channel in range(23755, 23826):
            dlfreq = 736.5 + (channel - 23755) * 0.1
            ulfreq = 706.5 + (channel - 23755) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 18':
        if channel in range(23875, 23976):
            dlfreq = 862.5 + (channel - 23875) * 0.1
            ulfreq = 817.5 + (channel - 23875) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 19':
        if channel in range(24025, 24126):
            dlfreq = 877.5 + (channel - 24025) * 0.1
            ulfreq = 832.5 + (channel - 24025) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 20':
        if channel in range(24175, 24426):
            dlfreq = 793.5 + (channel - 24175) * 0.1
            ulfreq = 834.5 + (channel - 24175) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 21':
        if channel in range(24475, 24576):
            dlfreq = 1498.4 + (channel - 24475) * 0.1
            ulfreq = 1450.4 + (channel - 24475) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 22':
        if channel in range(24625, 25376):
            dlfreq = 3512.5 + (channel - 24625) * 0.1
            ulfreq = 3412.5 + (channel - 24625) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 23':
        if channel in range(25507, 25694):
            dlfreq = 2180.7 + (channel - 25507) * 0.1
            ulfreq = 2000.7 + (channel - 25507) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 24':
        if channel in range(25725, 26016):
            dlfreq = 1527.5 + (channel - 25725) * 0.1
            ulfreq = 1629.0 + (channel - 25725) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 25':
        if channel in range(26047, 26684):
            dlfreq = 1930.7 + (channel - 26047) * 0.1
            ulfreq = 1850.7 + (channel - 26047) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 26':
        if channel in range(26697, 27034):
            dlfreq = 859.7 + (channel - 26697) * 0.1
            ulfreq = 814.7 + (channel - 26697) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 27':
        if channel in range(27047, 27204):
            dlfreq = 852.7 + (channel - 27047) * 0.1
            ulfreq = 807.7 + (channel - 27047) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 28':
        if channel in range(27225, 27646):
            dlfreq = 759.5 + (channel - 27225) * 0.1
            ulfreq = 704.5 + (channel - 27225) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 30':
        if channel in range(27685, 27736):
            dlfreq = 2352.5 + (channel - 27685) * 0.1
            ulfreq = 2307.5 + (channel - 27685) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'FDD-LTE 31':
        if channel in range(27767, 27804):
            dlfreq = 463.2 + (channel - 27767) * 0.1
            ulfreq = 453.2 + (channel - 27767) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'TDD-LTE 33':
        if channel in range(36025, 36176):
            dlfreq = 1902.5 + (channel - 36025) * 0.1
            ulfreq = 1902.5 + (channel - 36025) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'TDD-LTE 34':
        if channel in range(36225, 36326):
            dlfreq = 2012.5 + (channel - 36225) * 0.1
            ulfreq = 2012.5 + (channel - 36225) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'TDD-LTE 35':
        if channel in range(36357, 36944):
            dlfreq = 1850.7 + (channel - 36357) * 0.1
            ulfreq = 1850.7 + (channel - 36357) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'TDD-LTE 36':
        if channel in range(36957, 37544):
            dlfreq = 1930.7 + (channel - 36957) * 0.1
            ulfreq = 1930.7 + (channel - 36957) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'TDD-LTE 37':
        if channel in range(37575, 37726):
            dlfreq = 1912.5 + (channel - 37575) * 0.1
            ulfreq = 1912.5 + (channel - 37575) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band in ['TDD-LTE 38', 'CA_38C']:
        if channel in range(37775, 38226):
            dlfreq = 2572.5 + (channel - 37775) * 0.1
            ulfreq = 2572.5 + (channel - 37775) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'TDD-LTE 39':
        if channel in range(38275, 38626):
            dlfreq = 1882.5 + (channel - 38275) * 0.1
            ulfreq = 1882.5 + (channel - 38275) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'TDD-LTE 40':
        if channel in range(38675, 39626):
            dlfreq = 2302.5 + (channel - 38675) * 0.1
            ulfreq = 2302.5 + (channel - 38675) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band in ['TDD-LTE 41', 'CA_41C']:
        if channel in range(39675, 41566):
            dlfreq = 2498.5 + (channel - 39675) * 0.1
            ulfreq = 2498.5 + (channel - 39675) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    elif global_element.Test_band == 'TDD-LTE 42':
        if channel in range(41615, 43566):
            dlfreq = 3402.5 + (channel - 41615) * 0.1
            ulfreq = 3402.5 + (channel - 41615) * 0.1
        else:
            global_element.emitsingle.thread_exitSingle.emit('%s Channel definition error!' % global_element.Test_band)
    else:
        global_element.emitsingle.thread_exitSingle.emit('Conversion between channel and frequency is not yet defined '
                                                         'for %s' % global_element.Test_band)

    return str(format(dlfreq, '.1f')), str(format(ulfreq, '.1f'))


# 测试得到的数据输出到报告和界面
def Reporttool(result, lowlimit, highlimit):
    if global_element.IsStop == False:                                                    # 判断用户是否已经终止进程
        if not os.access(global_element.reportpath, os.F_OK):                             # 先判断report文件是否存存在， 不存在就新建
            reportwb = Workbook()
            sheets_list = reportwb.sheetnames
            if len(sheets_list) > 1:
                for i in range(1, len(sheets_list)):
                    reportwb.remove(reportwb[sheets_list[i]])   # 只保存一个sheet

            # reportwb.save
            reportwb.save(global_element.reportpath)
            reportwb.close()

        rpwb = load_workbook(global_element.reportpath)

        sheetname_list = rpwb.sheetnames
        if 'DUT Info' not in sheetname_list:
            rpwb.create_sheet('DUT Info')
            columns_list = ['DUT Name:', 'DUT Manufacturer:', 'DUT Serial Number:', 'DUT Hardware Revision:',
                            'DUT Software Revision:', 'DUT IMEI:', 'Maximum Registration Time(s):', 'Voltage High(V):',
                            'Voltage Normal(V):', 'Voltage Low(V):', 'Current Max(A):', 'Temperature High(℃):',
                            'Temperature Normal(℃):','Temperature Low(℃):']
            dut_config_dict = global_element.active_dut_dict['xml']['DUTCONFIG']
            DUT_value_list = [global_element.active_dut_name, dut_config_dict['DUTMANUFACTURER'], dut_config_dict['DUTSN'],
                              dut_config_dict['DUTHWREV'], dut_config_dict['DUTSWREV'], dut_config_dict['DUTIMEI'],
                              dut_config_dict['MAXREGTIME'], dut_config_dict['HV'], dut_config_dict['NV'],
                              dut_config_dict['LV'], dut_config_dict['MAXC'], dut_config_dict['HT'], dut_config_dict['NT'],
                              dut_config_dict['LT']]
            sheet = rpwb['DUT Info']
            for i in range(len(columns_list)):
                sheet.cell(row=i + 1, column=1).value = columns_list[i]
                sheet.cell(row=i + 1, column=2).value = DUT_value_list[i]

        # 如果是GSM测试，判断report中是否包含‘GSM’的工作表，没有就初始化一个。如果有就直接updata测试结果
        if global_element.Test_type == 'GSM' or global_element.Test_type == 'GSMFCC':
            sheetname_list = rpwb.sheetnames
            if 'GSM' not in sheetname_list:
                rpwb.create_sheet('GSM')
                list_head = ['Test Type', 'Test Case', 'Test Items', 'Result', 'Low Limit', 'High Limit', 'Judgement',
                             'Band', 'Channel', 'DL_Freq(MHz)', 'UL_Freq(MHz)', 'PCL', 'Modulation', 'Volt.(V)', 'Temp.(℃)',
                             'Remark', 'Time']
                sheet = rpwb['GSM']
                for i in range(len(list_head)):
                    sheet.cell(row=1, column=i+1).value = list_head[i]

            sheet_gsm = rpwb['GSM']
            Judgement = judgement_handle(result, lowlimit, highlimit)
            DL_Freq, UL_Freq = freq_gsm_calc(int(global_element.Test_channel))
            time_now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
            row_count = sheet_gsm.max_row
            list_data = [global_element.Test_type, global_element.Test_case, global_element.Test_item, result, lowlimit,
                         highlimit, Judgement, global_element.Test_band, global_element.Test_channel, DL_Freq, UL_Freq,
                         global_element.Test_pcl, global_element.Test_modulation, global_element.Test_volt,
                         global_element.Test_temp, global_element.Test_remark, time_now]
            for i in range(len(list_data)):
                sheet_gsm.cell(row=(row_count+1), column=i+1).value = list_data[i]
            global_element.emitsingle.reportupdataSingle.emit(list_data)
            for i in sheetname_list:
                if 'Sheet' in i:
                    rpwb.remove(rpwb[i])
                    break
            rpwb.save(global_element.reportpath)

            global_element.testseq_judgement_list.append(Judgement)

            if Judgement == 'Passed':
                global_element.finished_result_list[0] += 1
            elif Judgement == 'Failed':
                global_element.finished_result_list[1] += 1
            elif Judgement == 'Inconclusive':
                global_element.finished_result_list[2] += 1

        # 如果是WCDMA测试，判断report中是否包含‘WCDMA’的工作表，没有就初始化一个。如果有就直接updata测试结果
        if global_element.Test_type == 'WCDMA' or global_element.Test_type == 'WCDMAFCC':
            sheetname_list = rpwb.sheetnames
            if 'WCDMA' not in sheetname_list:
                rpwb.create_sheet('WCDMA')
                list_head = ['Test Type', 'Test Case', 'Test Items', 'Result', 'Low Limit', 'High Limit', 'Judgement',
                             'Band', 'Channel', 'DL_Freq(MHz)', 'UL_Freq(MHz)', 'Modulation', 'Volt.(V)', 'Temp.(℃)',
                             'Remark', 'Time']
                sheet = rpwb['WCDMA']
                for i in range(len(list_head)):
                    sheet.cell(row=1, column=i + 1).value = list_head[i]

            sheet_wcdma = rpwb['WCDMA']
            Judgement = judgement_handle(result, lowlimit, highlimit)
            DL_Freq, UL_Freq = freq_gsm_calc(int(global_element.Test_channel))
            time_now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
            row_count = sheet_wcdma.max_row
            list_data = [global_element.Test_type, global_element.Test_case, global_element.Test_item, result, lowlimit,
                         highlimit, Judgement, global_element.Test_band, global_element.Test_channel, DL_Freq, UL_Freq,
                         global_element.Test_modulation, global_element.Test_volt,
                         global_element.Test_temp, global_element.Test_remark, time_now]
            for i in range(len(list_data)):
                sheet_wcdma.cell(row=(row_count + 1), column=i + 1).value = list_data[i]
            global_element.emitsingle.reportupdataSingle.emit(list_data)
            for i in sheetname_list:
                if 'Sheet' in i:
                    rpwb.remove(rpwb[i])
                    break
            rpwb.save(global_element.reportpath)

            global_element.testseq_judgement_list.append(Judgement)

            if Judgement == 'Passed':
                global_element.finished_result_list[0] += 1
            elif Judgement == 'Failed':
                global_element.finished_result_list[1] += 1
            elif Judgement == 'Inconclusive':
                global_element.finished_result_list[2] += 1

        # 如果是LTE测试，判断report中是否包含‘LTE’的工作表，没有就初始化一个。如果有就直接updata测试结果
        if global_element.Test_type in ['LTE', 'LTEFCC']:
            sheetname_list = rpwb.sheetnames
            if 'LTE' not in sheetname_list:
                rpwb.create_sheet('LTE')
                list_head = ['Test Type', 'Test Case', 'Test Items', 'Result', 'Low Limit', 'High Limit', 'Judgement',
                          'Band', 'BandWidth', 'Channel', 'DL_Freq(MHz)', 'UL_Freq(MHz)', 'RB', 'Modulation', 'Volt.(V)',
                             'Temp.(℃)', 'Remark', 'Time']
                sheet = rpwb['LTE']
                for i in range(len(list_head)):
                    sheet.cell(row=1, column=i + 1).value = list_head[i]

            sheet_lte = rpwb['LTE']
            Judgement = judgement_handle(result, lowlimit, highlimit)
            DL_Freq, UL_Freq = freq_gsm_calc(int(global_element.Test_channel))
            time_now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
            row_count = sheet_lte.max_row
            list_data = [global_element.Test_type, global_element.Test_case, global_element.Test_item, result, lowlimit,
                         highlimit, Judgement, global_element.Test_band, global_element.Test_lte_bandwidth,
                         global_element.Test_channel, DL_Freq, UL_Freq,
                         global_element.Test_lte_RB, global_element.Test_modulation, global_element.Test_volt,
                         global_element.Test_temp, global_element.Test_remark, time_now]
            for i in range(len(list_data)):
                sheet_lte.cell(row=(row_count + 1), column=i + 1).value = list_data[i]
            global_element.emitsingle.reportupdataSingle.emit(list_data)
            for i in sheetname_list:
                if 'Sheet' in i:
                    rpwb.remove(rpwb[i])
                    break
            rpwb.save(global_element.reportpath)

            global_element.testseq_judgement_list.append(Judgement)

            if Judgement == 'Passed':
                global_element.finished_result_list[0] += 1
            elif Judgement == 'Failed':
                global_element.finished_result_list[1] += 1
            elif Judgement == 'Inconclusive':
                global_element.finished_result_list[2] += 1

        # LTE CA
        if global_element.Test_type in ['LTECA', 'LTECAFCC']:
            sheetname_list = rpwb.sheetnames
            if 'LTE_CA' not in sheetname_list:
                rpwb.create_sheet('LTE_CA')
                list_head = ['Test Type', 'Test Case', 'Test Items', 'Result', 'Low Limit', 'High Limit', 'Judgement',
                             'Band', 'BandWidth', 'Channel', 'DL_Freq(MHz)', 'UL_Freq(MHz)', 'RB', 'Modulation',
                             'Volt.(V)',
                             'Temp.(℃)', 'Remark', 'Time']
                sheet = rpwb['LTE_CA']
                for i in range(len(list_head)):
                    sheet.cell(row=1, column=i + 1).value = list_head[i]

            sheet_lte = rpwb['LTE_CA']
            Judgement = judgement_handle(result, lowlimit, highlimit)
            DL_Freq = ''
            UL_Freq = ''
            time_now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
            row_count = sheet_lte.max_row
            list_data = [global_element.Test_type, global_element.Test_case, global_element.Test_item, result, lowlimit,
                         highlimit, Judgement, global_element.Test_band, global_element.Test_lte_bandwidth,
                         global_element.Test_channel, DL_Freq, UL_Freq,
                         global_element.Test_lte_RB, global_element.Test_modulation, global_element.Test_volt,
                         global_element.Test_temp, global_element.Test_remark, time_now]
            for i in range(len(list_data)):
                sheet_lte.cell(row=(row_count + 1), column=i + 1).value = list_data[i]
            global_element.emitsingle.reportupdataSingle.emit(list_data)
            for i in sheetname_list:
                if 'Sheet' in i:
                    rpwb.remove(rpwb[i])
                    break
            rpwb.save(global_element.reportpath)

            global_element.testseq_judgement_list.append(Judgement)

            if Judgement == 'Passed':
                global_element.finished_result_list[0] += 1
            elif Judgement == 'Failed':
                global_element.finished_result_list[1] += 1
            elif Judgement == 'Inconclusive':
                global_element.finished_result_list[2] += 1

        # 如果是BT2测试，判断report中是否包含‘BT2’的工作表，没有就初始化一个。如果有就直接updata测试结果
        if global_element.Test_type == 'BT2' or global_element.Test_type == 'BT2FCC':
            sheetname_list = rpwb.sheetnames
            if 'BT2' not in sheetname_list:
                rpwb.create_sheet('BT2')
                list_head = ['Test Type', 'Test Case', 'Test Items', 'Result', 'Low Limit', 'High Limit',
                             'Judgement',
                             'Packet Type', 'Channel', 'Freq(MHz)',
                             'Volt.(V)',
                             'Temp.(℃)', 'Remark', 'Time']
                sheet = rpwb['BT2']
                for i in range(len(list_head)):
                    sheet.cell(row=1, column=i + 1).value = list_head[i]

            sheet_lte = rpwb['BT2']
            Judgement = judgement_handle(result, lowlimit, highlimit)
            Freq = fre_bt2_calc(global_element.Test_channel)
            time_now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
            row_count = sheet_lte.max_row
            list_data = [global_element.Test_type, global_element.Test_case, global_element.Test_item, result,
                         lowlimit,
                         highlimit, Judgement, global_element.Test_bt_packetype, global_element.Test_channel,
                         Freq,
                         global_element.Test_volt,
                         global_element.Test_temp, global_element.Test_remark, time_now]
            for i in range(len(list_data)):
                sheet_lte.cell(row=(row_count + 1), column=i + 1).value = list_data[i]
            global_element.emitsingle.reportupdataSingle.emit(list_data)
            for i in sheetname_list:
                if 'Sheet' in i:
                    rpwb.remove(rpwb[i])
                    break
            rpwb.save(global_element.reportpath)

            global_element.testseq_judgement_list.append(Judgement)

            if Judgement == 'Passed':
                global_element.finished_result_list[0] += 1
            elif Judgement == 'Failed':
                global_element.finished_result_list[1] += 1
            elif Judgement == 'Inconclusive':
                global_element.finished_result_list[2] += 1

        if global_element.Test_type == 'WLAN_AC':
            sheetname_list = rpwb.sheetnames
            if 'WLAN_AC' not in sheetname_list:
                rpwb.create_sheet('WLAN_AC')
                list_head = ['Test Type', 'Test Case', 'Test Items', 'Result', 'Low Limit', 'High Limit',
                             'Judgement',
                             'DataRate', 'BW', 'Channel', 'Freq(MHz)',
                             'Volt.(V)',
                             'Temp.(℃)', 'Remark', 'Time']
                sheet = rpwb['WLAN_AC']
                for i in range(len(list_head)):
                    sheet.cell(row=1, column=i + 1).value = list_head[i]

            sheet_lte = rpwb['WLAN_AC']
            Judgement = judgement_handle(result, lowlimit, highlimit)
            Freq = fre_wlan_calc(global_element.Test_channel)
            time_now = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
            row_count = sheet_lte.max_row
            list_data = [global_element.Test_type, global_element.Test_case, global_element.Test_item, result, lowlimit,
                         highlimit, Judgement, global_element.Test_Datarate, global_element.Test_wlan_bw,
                         global_element.Test_channel,
                         Freq,
                         global_element.Test_volt,
                         global_element.Test_temp, global_element.Test_remark, time_now]
            for i in range(len(list_data)):
                sheet_lte.cell(row=(row_count + 1), column=i + 1).value = list_data[i]
            global_element.emitsingle.reportupdataSingle.emit(list_data)
            for i in sheetname_list:
                if 'Sheet' in i:
                    rpwb.remove(rpwb[i])
                    break
            rpwb.save(global_element.reportpath)

            global_element.testseq_judgement_list.append(Judgement)

            if Judgement == 'Passed':
                global_element.finished_result_list[0] += 1
            elif Judgement == 'Failed':
                global_element.finished_result_list[1] += 1
            elif Judgement == 'Inconclusive':
                global_element.finished_result_list[2] += 1

        rpwb.close()


def tabletoreport(filepath, tableindex, expect_str):
    document = Document(filepath)
    expect_text = expect_str
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break

    tbl, p = tableindex._tbl, target._p
    p.addnext(tbl)
    document.save(filepath)


def fccreporthandle():
    global_element.emitsingle.stateupdataSingle.emit('Start sorting out the report!')
    # 判断是否有勾选gsm 或 wcdma
    if global_element.reporthandle_dict['checkgsm'] == 2 or global_element.reporthandle_dict['checkwcdma'] == 2:
        global_element.emitsingle.stateupdataSingle.emit('Start processing GSM & WCDMA!……')
        # 复制一个模板到用户定义的report路径
        report_path = global_element.reporthandle_dict['reportaddr']
        shutil.copy('Config/Report mask/GSM & WCDMA FCC report mask.docx', report_path)

        # 将用户定义的基本信息填入报告对应位置
        before_list = ['@inp_reportno', '@inp_projectname', '@inp_grantno', '@inp_fccid', '@inp_dateofreceipt',
                       '@inp_dateoftest', '@inp_dateofissue', '@inp_brand', '@inp_producttype', '@inp_addressofbrand',
                       '@inp_hwversion', '@inp_swversion', '@inp_testedbyname']
        after_list = [global_element.reporthandle_dict['reportno'], global_element.reporthandle_dict['objectname'],
                      global_element.reporthandle_dict['grantno'],
                      global_element.reporthandle_dict['grantno'] + global_element.reporthandle_dict['objectname'],
                      datezhuanhuan(global_element.reporthandle_dict['dateofrec']),
                      datezhuanhuan(global_element.reporthandle_dict['dateofstart']) + ' ~ ' +
                      datezhuanhuan(global_element.reporthandle_dict['dateofstop']),
                      datezhuanhuan(global_element.reporthandle_dict['dateofissue']),
                      global_element.reporthandle_dict['brand'],
                      global_element.reporthandle_dict['producttype'], global_element.reporthandle_dict['brandaddr'],
                      global_element.reporthandle_dict['hwversion'], global_element.reporthandle_dict['swversion'],
                      namezhuanhuan(global_element.reporthandle_dict['testedby'])]

        for index in range(len(before_list)):
            replaceword(report_path, before_list[index], after_list[index])

        # 处理gsm & wcdma conducted output power
        fccgsmwcdmapowertoword()

        # 处理gsm & wcdma Peak-To-Average Ratio
        fccgsmwcdmaptatoword()

        # 处理wcdma Peak-To-Average Ratio图片
        pic_name_list_wcdma = ['WCDMA Band II Channel9262', 'WCDMA Band II Channel9400', 'WCDMA Band II Channel9538',
                         'WCDMA Band IV Channel1312', 'WCDMA Band IV Channel1413', 'WCDMA Band IV Channel1513',
                         'WCDMA Band V Channel4132', 'WCDMA Band V Channel4183', 'WCDMA Band V Channel4233']
        fccgsmwcdma_pictoword('WCDMA FCC', 'Peak-to-Average Ratio', pic_name_list_wcdma, u'Peak-to-Average Ratio – Photograph')

        # 处理gsm Peak-To-Average Ratio图片
        pic_name_list_gsm = ['GSM GSM850 Channel128 PCL5', 'GSM GSM850 Channel190 PCL5', 'GSM GSM850 Channel251 PCL5',
                         'GSM PCS1900 Channel512 PCL0', 'GSM PCS1900 Channel661 PCL0', 'GSM PCS1900 Channel810 PCL0',
                         'GPRS GSM850 Channel128 PCL5', 'GPRS GSM850 Channel190 PCL5', 'GPRS GSM850 Channel251 PCL5',
                         'GPRS PCS1900 Channel512 PCL0', 'GPRS PCS1900 Channel661 PCL0', 'GPRS PCS1900 Channel810 PCL0',
                         'EGPRS GSM850 Channel128 PCL8', 'EGPRS GSM850 Channel190 PCL8', 'EGPRS GSM850 Channel251 PCL8',
                         'EGPRS PCS1900 Channel512 PCL2', 'EGPRS PCS1900 Channel661 PCL2',
                         'EGPRS PCS1900 Channel810 PCL2']
        fccgsmwcdma_pictoword('GSM FCC', 'Peak-to-Average Ratio', pic_name_list_gsm,
                              u'Peak-to-Average Ratio – Photograph')

        # gsm wcdma BW数据导入Word
        bw_data_to_word()

        # 处理wcdma 26db BW图片
        fccgsmwcdma_pictoword('WCDMA FCC', '26dB Occupied Bandwidth', pic_name_list_wcdma,
                              u'26dB Bandwidth – Photograph')

        # 处理gsm 26db BW图片
        fccgsmwcdma_pictoword('GSM FCC', '26dB Occupied Bandwidth', pic_name_list_gsm, u'26dB Bandwidth – Photograph')

        # 处理wcdma 99% BW图片
        fccgsmwcdma_pictoword('WCDMA FCC', '99% Occupied Bandwidth', pic_name_list_wcdma,
                              u'Occupied Bandwidth – Photograph')

        # 处理gsm 99% BW图片
        fccgsmwcdma_pictoword('GSM FCC', '99% Occupied Bandwidth', pic_name_list_gsm,
                              u'Occupied Bandwidth – Photograph')

        # gsm wcdma Bandedge数据导入Word
        fcc_gsmwcdma_bandedge_to_word()

        # 处理wcdma Bandedge图片
        pic_name_list_wcdma = ['WCDMA Band II Channel9262', 'WCDMA Band II Channel9538',
                               'WCDMA Band IV Channel1312', 'WCDMA Band IV Channel1513',
                               'WCDMA Band V Channel4132', 'WCDMA Band V Channel4233']
        fccgsmwcdma_pictoword('WCDMA FCC', 'Band edge', pic_name_list_wcdma,
                              u'Conducted Band Edge – Photograph')

        # 处理gsm Bandedge图片
        pic_name_list_gsm = ['GSM GSM850 Channel128 PCL5', 'GSM GSM850 Channel251 PCL5',
                             'GSM PCS1900 Channel512 PCL0', 'GSM PCS1900 Channel810 PCL0',
                             'GPRS GSM850 Channel128 PCL5', 'GPRS GSM850 Channel251 PCL5',
                             'GPRS PCS1900 Channel512 PCL0', 'GPRS PCS1900 Channel810 PCL0',
                             'EGPRS GSM850 Channel128 PCL8', 'EGPRS GSM850 Channel251 PCL8',
                             'EGPRS PCS1900 Channel512 PCL2', 'EGPRS PCS1900 Channel810 PCL2']
        fccgsmwcdma_pictoword('GSM FCC', 'Band edge', pic_name_list_gsm,
                              u'Conducted Band Edge – Photograph')

        # 处理GSM WCDMA CSE数据到word
        fcc_gsmwcdma_cse_to_word()

        pic_name_list_wcdma = ['WCDMA Band II Channel9262', 'WCDMA Band II Channel9400', 'WCDMA Band II Channel9538',
                         'WCDMA Band IV Channel1312', 'WCDMA Band IV Channel1413', 'WCDMA Band IV Channel1513',
                         'WCDMA Band V Channel4132', 'WCDMA Band V Channel4183', 'WCDMA Band V Channel4233']

        # 处理 WCDMA CSE图片到word
        fccgsmwcdma_pictoword('WCDMA FCC', 'Conducted Spurious emissions', pic_name_list_wcdma,
                              u'Conducted Spurious Emission – Photograph')

        pic_name_list_gsm = ['GSM GSM850 Channel128 PCL5', 'GSM GSM850 Channel190 PCL5', 'GSM GSM850 Channel251 PCL5',
                         'GSM PCS1900 Channel512 PCL0', 'GSM PCS1900 Channel661 PCL0', 'GSM PCS1900 Channel810 PCL0',
                         'GPRS GSM850 Channel128 PCL5', 'GPRS GSM850 Channel190 PCL5', 'GPRS GSM850 Channel251 PCL5',
                         'GPRS PCS1900 Channel512 PCL0', 'GPRS PCS1900 Channel661 PCL0', 'GPRS PCS1900 Channel810 PCL0',
                         'EGPRS GSM850 Channel128 PCL8', 'EGPRS GSM850 Channel190 PCL8', 'EGPRS GSM850 Channel251 PCL8',
                         'EGPRS PCS1900 Channel512 PCL2', 'EGPRS PCS1900 Channel661 PCL2',
                         'EGPRS PCS1900 Channel810 PCL2']

        # 处理 GSM CSE图片到word
        fccgsmwcdma_pictoword('GSM FCC', 'Conducted Spurious emissions', pic_name_list_gsm,
                              u'Conducted Spurious Emission – Photograph')

    global_element.emitsingle.stateupdataSingle.emit('Report generation is complete!')


def replaceword(filename, before, after):
    '''
        替换word中的指定内容
    :param filename: 为文件的绝对路径，不支持相对路径
    :param before:
    :param after:
    :return:
    '''
    word = Dispatch('Word.Application')
    word.Visible = 0
    word.DisplayAlerts = 0
    doc = word.Documents.Open(filename)
    word.Selection.Find.Execute(before, False, False, False, False, False, True, 1, True, after, 2)
    time.sleep(1)
    doc.Save()
    doc.Close()


# 在word指定位置插入空白行
def insert_kongbai_to_word(filepath, expect_str):
    doc = Document(filepath)
    tab = doc.add_table(rows=1, cols=1)

    expect_text = expect_str
    for paragraph in doc.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break

    tbl, p = tab._tbl, target._p
    p.addnext(tbl)
    doc.save(filepath)


# 将数字格式的日期转换为英文格式
def datezhuanhuan(date_str):
    date_list = date_str.split('/')
    year = date_list[0]
    if date_list[1] == '01':
        month = 'Jan'
    elif date_list[1] == '02':
        month = 'Feb'
    elif date_list[1] == '03':
        month = 'Mar'
    elif date_list[1] == '04':
        month = 'Apr'
    elif date_list[1] == '05':
        month = 'May'
    elif date_list[1] == '06':
        month = 'Jun'
    elif date_list[1] == '07':
        month = 'Jul'
    elif date_list[1] == '08':
        month = 'Aug'
    elif date_list[1] == '09':
        month = 'Sept'
    elif date_list[1] == '10':
        month = 'Oct'
    elif date_list[1] == '11':
        month = 'Nov'
    elif date_list[1] == '12':
        month = 'Dec'
    else:
        month = 'NULL'
    date = date_list[2]

    english_date_str = month + ' ' + date + ', ' + year
    return english_date_str


# 将中文名称改为英文
def namezhuanhuan(name_Chinese):
    if name_Chinese == '王亮':
        name_eng = 'Liang Wang'
    elif name_Chinese == '朱保华':
        name_eng = 'Baohua Zhu'
    elif name_Chinese == '鲁芬':
        name_eng = 'Fen Lu'
    elif name_Chinese == '陈燕':
        name_eng = 'Yan Chen'
    elif name_Chinese == '郑红艳':
        name_eng = 'Hongyan Zheng'
    elif name_Chinese == '李兴朗':
        name_eng = 'Xinglang Li'
    elif name_Chinese == '王莹':
        name_eng = 'Ying Wang'
    else:
        name_eng = 'NULL'

    return name_eng


# 处理gsm & wcdma conducted output power数据到word报告
def fccgsmwcdmapowertoword():
    global_element.emitsingle.stateupdataSingle.emit('正在处理GSM & WCDMA Conducted Output Power数据……')

    # 获取 GSM + WCDMA 的band总数，用以确定table的样式
    # datasource = global_element.reporthandle_dict['datasource']
    datasource = global_element.reporthandle_dict['datasource']
    datasource_wb = load_workbook(datasource)
    sheetname_list = datasource_wb.sheetnames
    gsm_band_list = []
    gprs_band_list = []
    egprs_band_list = []
    wcdma_band_list = []
    gsm_channel_list = []
    gprs_channel_list = []
    egprs_channel_list = []
    wcdma_channel_list = []
    gsm_band_count = 0
    wcdma_band_count = 0
    if 'GSM' in sheetname_list:
        sheet_gsm = datasource_wb['GSM']
        row_max_count = sheet_gsm.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_gsm.cell(row=row_count, column=8).value
            mode_value = sheet_gsm.cell(row=row_count, column=13).value
            channel_value = sheet_gsm.cell(row=row_count, column=9).value
            if band_value not in gsm_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GSM':
                gsm_band_list.append(band_value)
                if channel_value not in gsm_channel_list:
                    gsm_channel_list.append(channel_value)
            elif band_value not in gprs_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GPRS':
                gprs_band_list.append(band_value)
                if channel_value not in gprs_channel_list:
                    gprs_channel_list.append(channel_value)
            elif band_value not in egprs_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'EGPRS':
                egprs_band_list.append(band_value)
                if channel_value not in egprs_channel_list:
                    egprs_channel_list.append(channel_value)

        gsm_band_count = len(gsm_band_list)

    if 'WCDMA' in sheetname_list:
        sheet_wcdma = datasource_wb['WCDMA']
        row_max_count = sheet_wcdma.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_wcdma.cell(row=row_count, column=8).value
            channel_value = sheet_wcdma.cell(row=row_count, column=9).value
            if band_value not in wcdma_band_list and band_value in ['Band II', 'Band IV', 'Band V']:
                wcdma_band_list.append(band_value)
                if channel_value not in wcdma_channel_list:
                    wcdma_channel_list.append(channel_value)

        wcdma_band_count = len(wcdma_band_list)

    gms_wcdma_band_count = gsm_band_count + wcdma_band_count
    global_element.emitsingle.stateupdataSingle.emit('原始数据包含Band:' + '、'.join(gsm_band_list) + '、' +
                                                     '、'.join(wcdma_band_list))

    if len(egprs_band_list) > 0:
        # 提取出对应band数量的 EGPRS  power table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_power_need = tables[len(egprs_band_list) - 1]
        table_erp_need = tables[len(egprs_band_list) + 24]

        # EGPRS power数据填入表内
        band_jishu = 0
        if 'GSM850' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'EGPRS850'
            table_power_need.cell(row, coloumn - 1).text = '128'
            table_power_need.cell(row, coloumn).text = '190'
            table_power_need.cell(row, coloumn + 1).text = '251'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'EGPRS850'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '128'
            table_erp_need.cell(row_erp, coloumn_erp).text = '190'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128 = 'null'
            value_190 = 'null'
            value_251 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_EPR_128 = 'null'
            value_EPR_190 = 'null'
            value_EPR_251 = 'null'
            value_gain_128 = 'null'
            value_gain_190 = 'null'
            value_gain_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                gainvalue = sheet_gsm.cell(row=row_count, column=16).value
                judgementvalue = sheet_gsm.cell(row=row_count, column=7).value

                if bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '128' and \
                        itemvalue == 'Conducted output power':
                    value_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '128' and \
                        itemvalue == 'ERP':
                    value_EPR_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_128 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_128 = judgementvalue
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '190' and \
                        itemvalue == 'Conducted output power':
                    value_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '190' and \
                        itemvalue == 'ERP':
                    value_EPR_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_190 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_190 = judgementvalue
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '251' and \
                        itemvalue == 'Conducted output power':
                    value_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '251' and \
                        itemvalue == 'ERP':
                    value_EPR_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_251 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_251 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_power_need.cell(row + 1, coloumn).text = value_fre_190
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_power_need.cell(row + 2, coloumn - 1).text = value_128
            table_power_need.cell(row + 2, coloumn).text = value_190
            table_power_need.cell(row + 2, coloumn + 1).text = value_251

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_128
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_190
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_251
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_128
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_190
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_251
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_128
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_190
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_251
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'ERP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 7W(38.45dBm)'

            if all([value_judgement_128 == 'Passed', value_judgement_190 == 'Passed', value_judgement_251 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  Conducted Output Power没有获取有效的 EGPRS Channel 128的值！')
            if value_fre_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  Conducted Output Power没有获取有效的 EGPRS Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  Conducted Output Power没有获取有效的 EGPRS Channel 251的值！')

        if 'PCS1900' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'EGPRS1900'
            table_power_need.cell(row, coloumn - 1).text = '512'
            table_power_need.cell(row, coloumn).text = '661'
            table_power_need.cell(row, coloumn + 1).text = '810'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'EGPRS1900'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '512'
            table_erp_need.cell(row_erp, coloumn_erp).text = '661'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512 = 'null'
            value_661 = 'null'
            value_810 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_EPR_512 = 'null'
            value_EPR_661 = 'null'
            value_EPR_810 = 'null'
            value_gain_512 = 'null'
            value_gain_661 = 'null'
            value_gain_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                gainvalue = sheet_gsm.cell(row=row_count, column=16).value
                judgementvalue = sheet_gsm.cell(row=row_count, column=7).value

                if bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '512' and \
                        itemvalue == 'Conducted output power':
                    value_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '512' and \
                        itemvalue == 'EIRP':
                    value_EPR_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_512 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_512 = judgementvalue
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '661' and \
                        itemvalue == 'Conducted output power':
                    value_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '661' and \
                        itemvalue == 'EIRP':
                    value_EPR_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_661 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_661 = judgementvalue
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '810' and \
                        itemvalue == 'Conducted output power':
                    value_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '810' and \
                        itemvalue == 'EIRP':
                    value_EPR_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_810 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_810 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_power_need.cell(row + 1, coloumn).text = value_fre_661
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_power_need.cell(row + 2, coloumn - 1).text = value_512
            table_power_need.cell(row + 2, coloumn).text = value_661
            table_power_need.cell(row + 2, coloumn + 1).text = value_810

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_512
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_661
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_810
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_512
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_661
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_810
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_512
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_661
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_810
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'EIRP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 2W(33dBm)'

            if all([value_judgement_512 == 'Passed', value_judgement_661 == 'Passed', value_judgement_810 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 EGPRS Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 EGPRS Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 EGPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_power_need,
                      u'Conducted Output Power (Average power) – Result')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_erp_need, u'EIPR / ERP Result')

    if len(gprs_band_list) > 0:
        # 提取出对应band数量的 GPRS  power table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_power_need = tables[len(gprs_band_list) - 1]
        table_erp_need = tables[len(gprs_band_list) + 24]

        # GPRS power数据填入表内
        band_jishu = 0
        if 'GSM850' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'GPRS850'
            table_power_need.cell(row, coloumn - 1).text = '128'
            table_power_need.cell(row, coloumn).text = '190'
            table_power_need.cell(row, coloumn + 1).text = '251'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'GPRS850'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '128'
            table_erp_need.cell(row_erp, coloumn_erp).text = '190'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128 = 'null'
            value_190 = 'null'
            value_251 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_EPR_128 = 'null'
            value_EPR_190 = 'null'
            value_EPR_251 = 'null'
            value_gain_128 = 'null'
            value_gain_190 = 'null'
            value_gain_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                gainvalue = sheet_gsm.cell(row=row_count, column=16).value
                judgementvalue = sheet_gsm.cell(row=row_count, column=7).value

                if bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '128' and \
                        itemvalue == 'Conducted output power':
                    value_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '128' and \
                        itemvalue == 'ERP':
                    value_EPR_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_128 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_128 = judgementvalue
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '190' and \
                        itemvalue == 'Conducted output power':
                    value_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '190' and \
                        itemvalue == 'ERP':
                    value_EPR_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_190 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_190 = judgementvalue
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '251' and \
                        itemvalue == 'Conducted output power':
                    value_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '251' and \
                        itemvalue == 'ERP':
                    value_EPR_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_251 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_251 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_power_need.cell(row + 1, coloumn).text = value_fre_190
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_power_need.cell(row + 2, coloumn - 1).text = value_128
            table_power_need.cell(row + 2, coloumn).text = value_190
            table_power_need.cell(row + 2, coloumn + 1).text = value_251

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_128
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_190
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_251
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_128
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_190
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_251
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_128
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_190
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_251
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'ERP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 7W(38.45dBm)'

            if all([value_judgement_128 == 'Passed', value_judgement_190 == 'Passed', value_judgement_251 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GPRS Channel 128的值！')
            if value_fre_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GPRS Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GPRS Channel 251的值！')

        if 'PCS1900' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'GPRS1900'
            table_power_need.cell(row, coloumn - 1).text = '512'
            table_power_need.cell(row, coloumn).text = '661'
            table_power_need.cell(row, coloumn + 1).text = '810'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'GPRS1900'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '512'
            table_erp_need.cell(row_erp, coloumn_erp).text = '661'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512 = 'null'
            value_661 = 'null'
            value_810 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_EPR_512 = 'null'
            value_EPR_661 = 'null'
            value_EPR_810 = 'null'
            value_gain_512 = 'null'
            value_gain_661 = 'null'
            value_gain_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                gainvalue = sheet_gsm.cell(row=row_count, column=16).value
                judgementvalue = sheet_gsm.cell(row=row_count, column=7).value

                if bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '512' and \
                        itemvalue == 'Conducted output power':
                    value_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '512' and \
                        itemvalue == 'EIRP':
                    value_EPR_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_512 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_512 = judgementvalue
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '661' and \
                        itemvalue == 'Conducted output power':
                    value_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '661' and \
                        itemvalue == 'EIRP':
                    value_EPR_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_661 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_661 = judgementvalue
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '810' and \
                        itemvalue == 'Conducted output power':
                    value_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '810' and \
                        itemvalue == 'EIRP':
                    value_EPR_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_810 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_810 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_power_need.cell(row + 1, coloumn).text = value_fre_661
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_power_need.cell(row + 2, coloumn - 1).text = value_512
            table_power_need.cell(row + 2, coloumn).text = value_661
            table_power_need.cell(row + 2, coloumn + 1).text = value_810

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_512
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_661
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_810
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_512
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_661
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_810
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_512
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_661
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_810
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'EIRP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 2W(33dBm)'

            if all([value_judgement_512 == 'Passed', value_judgement_661 == 'Passed', value_judgement_810 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GPRS Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GPRS Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_power_need,
                      u'Conducted Output Power (Average power) – Result')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_erp_need, u'EIPR / ERP Result')

    if gms_wcdma_band_count > 0:
        # 提取出对应band数量的 GSM + WCDMA  power table模板

        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_power_need = tables[len(gsm_band_list) + len(wcdma_band_list) - 1]
        table_erp_need = tables[len(gsm_band_list) + len(wcdma_band_list) + 24]

        # GSM power数据填入表内
        band_jishu = 0
        if 'GSM850' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'GSM850'
            table_power_need.cell(row, coloumn - 1).text = '128'
            table_power_need.cell(row, coloumn).text = '190'
            table_power_need.cell(row, coloumn + 1).text = '251'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'GSM850'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '128'
            table_erp_need.cell(row_erp, coloumn_erp).text = '190'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128 = 'null'
            value_190 = 'null'
            value_251 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_EPR_128 = 'null'
            value_EPR_190 = 'null'
            value_EPR_251 = 'null'
            value_gain_128 = 'null'
            value_gain_190 = 'null'
            value_gain_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                gainvalue = sheet_gsm.cell(row=row_count, column=16).value
                judgementvalue = sheet_gsm.cell(row=row_count, column=7).value

                if bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '128' and \
                        itemvalue == 'Conducted output power':
                    value_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '128' and \
                        itemvalue == 'ERP':
                    value_EPR_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_128 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_128 = judgementvalue
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '190' and \
                        itemvalue == 'Conducted output power':
                    value_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '190' and \
                        itemvalue == 'ERP':
                    value_EPR_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_190 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_190 = judgementvalue
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '251' and \
                        itemvalue == 'Conducted output power':
                    value_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '251' and \
                        itemvalue == 'ERP':
                    value_EPR_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_251 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_251 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_power_need.cell(row + 1, coloumn).text = value_fre_190
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_power_need.cell(row + 2, coloumn - 1).text = value_128
            table_power_need.cell(row + 2, coloumn).text = value_190
            table_power_need.cell(row + 2, coloumn + 1).text = value_251

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_128
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_190
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_251
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_128
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_190
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_251
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_128
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_190
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_251
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'ERP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 7W(38.45dBm)'

            if all([value_judgement_128 == 'Passed', value_judgement_190 == 'Passed', value_judgement_251 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GSM Channel 128的值！')
            if value_fre_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GSM Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GSM Channel 251的值！')

        if 'PCS1900' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'PCS1900'
            table_power_need.cell(row, coloumn - 1).text = '512'
            table_power_need.cell(row, coloumn).text = '661'
            table_power_need.cell(row, coloumn + 1).text = '810'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'PCS1900'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '512'
            table_erp_need.cell(row_erp, coloumn_erp).text = '661'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512 = 'null'
            value_661 = 'null'
            value_810 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_EPR_512 = 'null'
            value_EPR_661 = 'null'
            value_EPR_810 = 'null'
            value_gain_512 = 'null'
            value_gain_661 = 'null'
            value_gain_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                gainvalue = sheet_gsm.cell(row=row_count, column=16).value
                judgementvalue = sheet_gsm.cell(row=row_count, column=7).value

                if bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '512' and \
                        itemvalue == 'Conducted output power':
                    value_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '512' and \
                        itemvalue == 'EIRP':
                    value_EPR_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_512 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_512 = judgementvalue
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '661' and \
                        itemvalue == 'Conducted output power':
                    value_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '661' and \
                        itemvalue == 'EIRP':
                    value_EPR_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_661 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_661 = judgementvalue
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '810' and \
                        itemvalue == 'Conducted output power':
                    value_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '810' and \
                        itemvalue == 'EIRP':
                    value_EPR_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_gain_810 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_810 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_power_need.cell(row + 1, coloumn).text = value_fre_661
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_power_need.cell(row + 2, coloumn - 1).text = value_512
            table_power_need.cell(row + 2, coloumn).text = value_661
            table_power_need.cell(row + 2, coloumn + 1).text = value_810

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_512
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_661
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_810
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_512
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_661
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_810
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_512
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_661
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_810
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'EIRP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 2W(33dBm)'

            if all([value_judgement_512 == 'Passed', value_judgement_661 == 'Passed', value_judgement_810 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GSM Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GSM Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 GSM Channel 810的值！')

        if 'Band II' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'Band II'
            table_power_need.cell(row, coloumn - 1).text = '9262'
            table_power_need.cell(row, coloumn).text = '9400'
            table_power_need.cell(row, coloumn + 1).text = '9538'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'Band II'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '9262'
            table_erp_need.cell(row_erp, coloumn_erp).text = '9400'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '9538'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_9262 = 'null'
            value_9400 = 'null'
            value_9538 = 'null'
            value_fre_9262 = 'null'
            value_fre_9400 = 'null'
            value_fre_9538 = 'null'
            value_EPR_9262 = 'null'
            value_EPR_9400 = 'null'
            value_EPR_9538 = 'null'
            value_gain_9262 = 'null'
            value_gain_9400 = 'null'
            value_gain_9538 = 'null'
            value_judgement_9262 = 'null'
            value_judgement_9400 = 'null'
            value_judgement_9538 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                gainvalue = sheet_wcdma.cell(row=row_count, column=15).value
                judgementvalue = sheet_wcdma.cell(row=row_count, column=7).value

                if bandvalue == 'Band II' and channelvalue == '9262' and \
                        itemvalue == 'Conducted output power':
                    value_9262 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9262 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band II' and channelvalue == '9262' and \
                        itemvalue == 'EIRP':
                    value_EPR_9262 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_9262 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_9262 = judgementvalue
                elif bandvalue == 'Band II' and channelvalue == '9400' and \
                        itemvalue == 'Conducted output power':
                    value_9400 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9400 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band II' and channelvalue == '9400' and \
                        itemvalue == 'EIRP':
                    value_EPR_9400 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_9400 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_9400 = judgementvalue
                elif bandvalue == 'Band II' and channelvalue == '9538' and \
                        itemvalue == 'Conducted output power':
                    value_9538 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9538 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band II' and channelvalue == '9538' and \
                        itemvalue == 'EIRP':
                    value_EPR_9538 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_9538 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_9538 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_9262
            table_power_need.cell(row + 1, coloumn).text = value_fre_9400
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_9538
            table_power_need.cell(row + 2, coloumn - 1).text = value_9262
            table_power_need.cell(row + 2, coloumn).text = value_9400
            table_power_need.cell(row + 2, coloumn + 1).text = value_9538

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_9262
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_9400
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_9538
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_9262
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_9400
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_9538
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_9262
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_9400
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_9538
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'EIRP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 2W(33dBm)'

            if all([value_judgement_9262 == 'Passed', value_judgement_9400 == 'Passed', value_judgement_9538 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_9262 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 9262的值！')
            if value_fre_9400 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 9400的值！')
            if value_fre_9538 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 9538的值！')

        if 'Band IV' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'Band IV'
            table_power_need.cell(row, coloumn - 1).text = '1312'
            table_power_need.cell(row, coloumn).text = '1413'
            table_power_need.cell(row, coloumn + 1).text = '1513'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'Band IV'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '1312'
            table_erp_need.cell(row_erp, coloumn_erp).text = '1413'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '1513'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_1312 = 'null'
            value_1413 = 'null'
            value_1513 = 'null'
            value_fre_1312 = 'null'
            value_fre_1413 = 'null'
            value_fre_1513 = 'null'
            value_EPR_1312 = 'null'
            value_EPR_1413 = 'null'
            value_EPR_1513 = 'null'
            value_gain_1312 = 'null'
            value_gain_1413 = 'null'
            value_gain_1513 = 'null'
            value_judgement_1312 = 'null'
            value_judgement_1413 = 'null'
            value_judgement_1513 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                gainvalue = sheet_wcdma.cell(row=row_count, column=15).value
                judgementvalue = sheet_wcdma.cell(row=row_count, column=7).value

                if bandvalue == 'Band IV' and channelvalue == '1312' and \
                        itemvalue == 'Conducted output power':
                    value_1312 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1312 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band IV' and channelvalue == '1312' and \
                        itemvalue == 'EIRP':
                    value_EPR_1312 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_1312 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_1312 = judgementvalue
                elif bandvalue == 'Band IV' and channelvalue == '1413' and \
                        itemvalue == 'Conducted output power':
                    value_1413 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1413 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band IV' and channelvalue == '1413' and \
                        itemvalue == 'EIRP':
                    value_EPR_1413 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_1413 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_1413 = judgementvalue
                elif bandvalue == 'Band IV' and channelvalue == '1513' and \
                        itemvalue == 'Conducted output power':
                    value_1513 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1513 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band IV' and channelvalue == '1513' and \
                        itemvalue == 'EIRP':
                    value_EPR_1513 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_1513 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_1513 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_1312
            table_power_need.cell(row + 1, coloumn).text = value_fre_1413
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_1513
            table_power_need.cell(row + 2, coloumn - 1).text = value_1312
            table_power_need.cell(row + 2, coloumn).text = value_1413
            table_power_need.cell(row + 2, coloumn + 1).text = value_1513

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_1312
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_1413
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_1513
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_1312
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_1413
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_1513
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_1312
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_1413
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_1513
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'EIRP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 1W(30dBm)'

            if all([value_judgement_1312 == 'Passed', value_judgement_1413 == 'Passed', value_judgement_1513 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_1312 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 1312的值！')
            if value_fre_1413 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 1413的值！')
            if value_fre_1513 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 1513的值！')

        if 'Band V' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 4 + 1
            coloumn = ((band_jishu - 1) % 2) * 4 + 2
            row_erp = ((band_jishu - 1) // 2) * 7 + 1
            coloumn_erp = ((band_jishu - 1) % 2) * 4 + 2
            table_power_need.cell(row - 1, coloumn - 1).text = 'Band V'
            table_power_need.cell(row, coloumn - 1).text = '4132'
            table_power_need.cell(row, coloumn).text = '4183'
            table_power_need.cell(row, coloumn + 1).text = '4233'
            table_erp_need.cell(row_erp - 1, coloumn_erp - 1).text = 'Band V'
            table_erp_need.cell(row_erp, coloumn_erp - 1).text = '4132'
            table_erp_need.cell(row_erp, coloumn_erp).text = '4183'
            table_erp_need.cell(row_erp, coloumn_erp + 1).text = '4233'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_4132 = 'null'
            value_4183 = 'null'
            value_4233 = 'null'
            value_fre_4132 = 'null'
            value_fre_4183 = 'null'
            value_fre_4233 = 'null'
            value_EPR_4132 = 'null'
            value_EPR_4183 = 'null'
            value_EPR_4233 = 'null'
            value_gain_4132 = 'null'
            value_gain_4183 = 'null'
            value_gain_4233 = 'null'
            value_judgement_4132 = 'null'
            value_judgement_4183 = 'null'
            value_judgement_4233 = 'null'

            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                gainvalue = sheet_wcdma.cell(row=row_count, column=15).value
                judgementvalue = sheet_wcdma.cell(row=row_count, column=7).value

                if bandvalue == 'Band V' and channelvalue == '4132' and \
                        itemvalue == 'Conducted output power':
                    value_4132 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4132 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band V' and channelvalue == '4132' and \
                        itemvalue == 'ERP':
                    value_EPR_4132 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_4132 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_4132 = judgementvalue
                elif bandvalue == 'Band V' and channelvalue == '4183' and \
                        itemvalue == 'Conducted output power':
                    value_4183 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4183 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band V' and channelvalue == '4183' and \
                        itemvalue == 'ERP':
                    value_EPR_4183 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_4183 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_4183 = judgementvalue
                elif bandvalue == 'Band V' and channelvalue == '4233' and \
                        itemvalue == 'Conducted output power':
                    value_4233 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4233 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band V' and channelvalue == '4233' and \
                        itemvalue == 'ERP':
                    value_EPR_4233 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_gain_4233 = gainvalue.split('(')[1].split(')')[0]
                    value_judgement_4233 = judgementvalue

            table_power_need.cell(row + 1, coloumn - 1).text = value_fre_4132
            table_power_need.cell(row + 1, coloumn).text = value_fre_4183
            table_power_need.cell(row + 1, coloumn + 1).text = value_fre_4233
            table_power_need.cell(row + 2, coloumn - 1).text = value_4132
            table_power_need.cell(row + 2, coloumn).text = value_4183
            table_power_need.cell(row + 2, coloumn + 1).text = value_4233

            table_erp_need.cell(row_erp + 1, coloumn_erp - 1).text = value_fre_4132
            table_erp_need.cell(row_erp + 1, coloumn_erp).text = value_fre_4183
            table_erp_need.cell(row_erp + 1, coloumn_erp + 1).text = value_fre_4233
            table_erp_need.cell(row_erp + 2, coloumn_erp - 1).text = value_gain_4132
            table_erp_need.cell(row_erp + 2, coloumn_erp).text = value_gain_4183
            table_erp_need.cell(row_erp + 2, coloumn_erp + 1).text = value_gain_4233
            table_erp_need.cell(row_erp + 3, coloumn_erp - 1).text = value_EPR_4132
            table_erp_need.cell(row_erp + 3, coloumn_erp).text = value_EPR_4183
            table_erp_need.cell(row_erp + 3, coloumn_erp + 1).text = value_EPR_4233
            table_erp_need.cell(row_erp + 3, coloumn_erp - 2).text = 'ERP(dBm)'
            table_erp_need.cell(row_erp + 4, coloumn_erp - 1).text = '<= 7W(38.45dBm)'

            if all([value_judgement_4132 == 'Passed', value_judgement_4183 == 'Passed', value_judgement_4233 == 'Passed']):
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Passed'
            else:
                table_erp_need.cell(row_erp + 5, coloumn_erp - 1).text = 'Failed'
                global_element.fcc_gsm_wcdma_erp_judgement = 'Failed'

            if value_fre_4132 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 4132的值！')
            if value_fre_4183 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 4183的值！')
            if value_fre_4233 == 'null':
                global_element.emitsingle.stateupdataSingle.emit(
                    '错误：  Conducted Output Power没有获取有效的 WCDMA Channel 4233的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_power_need,
                      u'Conducted Output Power (Average power) – Result')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_erp_need, u'EIPR / ERP Result')


# 处理gsm & wcdma Peak-To-Average Ratio数据到word报告
def fccgsmwcdmaptatoword():
    global_element.emitsingle.stateupdataSingle.emit('正在处理GSM & WCDMA Peak-To-Average Ratio数据……')

    # 获取 GSM + WCDMA 的band总数，用以确定table的样式
    datasource = global_element.reporthandle_dict['datasource']
    datasource_wb = load_workbook(datasource)
    sheetname_list = datasource_wb.sheetnames
    gsm_band_list = []
    gprs_band_list = []
    egprs_band_list = []
    wcdma_band_list = []
    gsm_band_count = 0
    wcdma_band_count = 0
    if 'GSM' in sheetname_list:
        sheet_gsm = datasource_wb['GSM']
        row_max_count = sheet_gsm.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_gsm.cell(row=row_count, column=8).value
            mode_value = sheet_gsm.cell(row=row_count, column=13).value
            item_value = sheet_gsm.cell(row=row_count, column=3).value
            if band_value not in gsm_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GSM' and \
                    item_value == 'Peak-to-Average Ratio':
                gsm_band_list.append(band_value)
            elif band_value not in gprs_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GPRS' and \
                    item_value == 'Peak-to-Average Ratio':
                gprs_band_list.append(band_value)
            elif band_value not in egprs_band_list and band_value in ['GSM850', 'PCS1900'] and \
                    mode_value == 'EGPRS' and item_value == 'Peak-to-Average Ratio':
                egprs_band_list.append(band_value)

        gsm_band_count = len(gsm_band_list)

    if 'WCDMA' in sheetname_list:
        sheet_wcdma = datasource_wb['WCDMA']
        row_max_count = sheet_wcdma.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_wcdma.cell(row=row_count, column=8).value
            item_value = sheet_wcdma.cell(row=row_count, column=3).value
            if band_value not in wcdma_band_list and band_value in ['Band II', 'Band IV', 'Band V'] and \
                    item_value == 'Peak-to-Average Ratio':
                wcdma_band_list.append(band_value)

        wcdma_band_count = len(wcdma_band_list)

    gms_wcdma_band_count = gsm_band_count + wcdma_band_count
    global_element.emitsingle.stateupdataSingle.emit('原始数据包含Band:' + '、'.join(gsm_band_list) + '、' +
                                                     '、'.join(wcdma_band_list))

    if len(egprs_band_list) > 0:
        # 提取出对应band数量的 EGPRS  peak to average table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(egprs_band_list) + 4]

        # EGPRS 数据填入表内
        band_jishu = 0
        if 'GSM850' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'EGPRS850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128 = 'null'
            value_190 = 'null'
            value_251 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '128' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '190' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_190 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '251' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_190, value_judgement_251]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_need.cell(row + 1, coloumn).text = value_fre_190
            table_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_need.cell(row + 2, coloumn - 1).text = value_128
            table_need.cell(row + 2, coloumn).text = value_190
            table_need.cell(row + 2, coloumn + 1).text = value_251
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 3, coloumn).text = value_judgement_190
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_251

            if value_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 128的值！')
            if value_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 190的值！')
            if value_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 251的值！')

        if 'PCS1900' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'EGPRS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512 = 'null'
            value_661 = 'null'
            value_810 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '512' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '661' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_661 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '810' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_661, value_judgement_810]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_need.cell(row + 1, coloumn).text = value_fre_661
            table_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_need.cell(row + 2, coloumn - 1).text = value_512
            table_need.cell(row + 2, coloumn).text = value_661
            table_need.cell(row + 2, coloumn + 1).text = value_810
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 3, coloumn).text = value_judgement_661
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_810

            if value_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 512的值！')
            if value_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 661的值！')
            if value_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need, u'Peak-to-Average Ratio – Result')

    if len(gprs_band_list) > 0:
        # 提取出对应band数量的 GPRS  peak to average table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(egprs_band_list) + 4]

        # GPRS 数据填入表内
        band_jishu = 0
        if 'GSM850' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GPRS850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128 = 'null'
            value_190 = 'null'
            value_251 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '128' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '190' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_190 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '251' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_190, value_judgement_251]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_need.cell(row + 1, coloumn).text = value_fre_190
            table_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_need.cell(row + 2, coloumn - 1).text = value_128
            table_need.cell(row + 2, coloumn).text = value_190
            table_need.cell(row + 2, coloumn + 1).text = value_251
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 3, coloumn).text = value_judgement_190
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_251

            if value_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 128的值！')
            if value_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 190的值！')
            if value_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 251的值！')

        if 'PCS1900' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GPRS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512 = 'null'
            value_661 = 'null'
            value_810 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '512' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '661' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_661 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '810' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_661, value_judgement_810]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_need.cell(row + 1, coloumn).text = value_fre_661
            table_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_need.cell(row + 2, coloumn - 1).text = value_512
            table_need.cell(row + 2, coloumn).text = value_661
            table_need.cell(row + 2, coloumn + 1).text = value_810
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 3, coloumn).text = value_judgement_661
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_810

            if value_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 512的值！')
            if value_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 661的值！')
            if value_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need, u'Peak-to-Average Ratio – Result')

    if gms_wcdma_band_count > 0:
        # 提取出对应band数量的 gsm wcdma  peak to average table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[gms_wcdma_band_count + 4]

        # GSM WCDMA 数据填入表内
        band_jishu = 0
        if 'GSM850' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GSM850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128 = 'null'
            value_190 = 'null'
            value_251 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '128' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_128 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '190' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_190 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_190 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '251' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_251 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_190, value_judgement_251]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_need.cell(row + 1, coloumn).text = value_fre_190
            table_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_need.cell(row + 2, coloumn - 1).text = value_128
            table_need.cell(row + 2, coloumn).text = value_190
            table_need.cell(row + 2, coloumn + 1).text = value_251
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 3, coloumn).text = value_judgement_190
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_251

            if value_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 128的值！')
            if value_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 190的值！')
            if value_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 251的值！')

        if 'PCS1900' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'PCS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512 = 'null'
            value_661 = 'null'
            value_810 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '512' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_512 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '661' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_661 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_661 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '810' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_810 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_661, value_judgement_810]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_need.cell(row + 1, coloumn).text = value_fre_661
            table_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_need.cell(row + 2, coloumn - 1).text = value_512
            table_need.cell(row + 2, coloumn).text = value_661
            table_need.cell(row + 2, coloumn + 1).text = value_810
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 3, coloumn).text = value_judgement_661
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_810

            if value_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 512的值！')
            if value_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 661的值！')
            if value_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 810的值！')

        if 'Band II' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band II'
            table_need.cell(row, coloumn - 1).text = '9262'
            table_need.cell(row, coloumn).text = '9400'
            table_need.cell(row, coloumn + 1).text = '9538'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_9262 = 'null'
            value_9400 = 'null'
            value_9538 = 'null'
            value_fre_9262 = 'null'
            value_fre_9400 = 'null'
            value_fre_9538 = 'null'
            value_judgement_9262 = 'null'
            value_judgement_9400 = 'null'
            value_judgement_9538 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band II' and channelvalue == '9262' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_9262 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9262 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_9262 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band II' and channelvalue == '9400' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_9400 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9400 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_9400 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band II' and channelvalue == '9538' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_9538 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9538 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_9538 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_9262, value_judgement_9400, value_judgement_9538]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_9262
            table_need.cell(row + 1, coloumn).text = value_fre_9400
            table_need.cell(row + 1, coloumn + 1).text = value_fre_9538
            table_need.cell(row + 2, coloumn - 1).text = value_9262
            table_need.cell(row + 2, coloumn).text = value_9400
            table_need.cell(row + 2, coloumn + 1).text = value_9538
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_9262
            table_need.cell(row + 3, coloumn).text = value_judgement_9400
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_9538

            if value_9262 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 9262的值！')
            if value_9400 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 9400的值！')
            if value_9538 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 9538的值！')

        if 'Band IV' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band IV'
            table_need.cell(row, coloumn - 1).text = '1312'
            table_need.cell(row, coloumn).text = '1413'
            table_need.cell(row, coloumn + 1).text = '1513'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_1312 = 'null'
            value_1413 = 'null'
            value_1513 = 'null'
            value_fre_1312 = 'null'
            value_fre_1413 = 'null'
            value_fre_1513 = 'null'
            value_judgement_1312 = 'null'
            value_judgement_1413 = 'null'
            value_judgement_1513 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band IV' and channelvalue == '1312' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_1312 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1312 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_1312 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band IV' and channelvalue == '1413' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_1413 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1413 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_1413 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band IV' and channelvalue == '1513' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_1513 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1513 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_1513 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_1312, value_judgement_1413, value_judgement_1513]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_1312
            table_need.cell(row + 1, coloumn).text = value_fre_1413
            table_need.cell(row + 1, coloumn + 1).text = value_fre_1513
            table_need.cell(row + 2, coloumn - 1).text = value_1312
            table_need.cell(row + 2, coloumn).text = value_1413
            table_need.cell(row + 2, coloumn + 1).text = value_1513
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_1312
            table_need.cell(row + 3, coloumn).text = value_judgement_1413
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_1513

            if value_1312 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 1312的值！')
            if value_1413 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 1413的值！')
            if value_1513 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 1513的值！')

        if 'Band V' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band V'
            table_need.cell(row, coloumn - 1).text = '4132'
            table_need.cell(row, coloumn).text = '4183'
            table_need.cell(row, coloumn + 1).text = '4233'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_4132 = 'null'
            value_4183 = 'null'
            value_4233 = 'null'
            value_fre_4132 = 'null'
            value_fre_4183 = 'null'
            value_fre_4233 = 'null'
            value_judgement_4132 = 'null'
            value_judgement_4183 = 'null'
            value_judgement_4233 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band V' and channelvalue == '4132' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_4132 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4132 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_4132 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band V' and channelvalue == '4183' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_4183 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4183 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_4183 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band V' and channelvalue == '4233' and \
                        itemvalue == 'Peak-to-Average Ratio':
                    value_4233 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4233 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_4233 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_4132, value_judgement_4183, value_judgement_4233]:
                global_element.fcc_gsm_wcdma_pta_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_fre_4132
            table_need.cell(row + 1, coloumn).text = value_fre_4183
            table_need.cell(row + 1, coloumn + 1).text = value_fre_4233
            table_need.cell(row + 2, coloumn - 1).text = value_4132
            table_need.cell(row + 2, coloumn).text = value_4183
            table_need.cell(row + 2, coloumn + 1).text = value_4233
            table_need.cell(row + 3, coloumn - 1).text = value_judgement_4132
            table_need.cell(row + 3, coloumn).text = value_judgement_4183
            table_need.cell(row + 3, coloumn + 1).text = value_judgement_4233

            if value_4132 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 4132的值！')
            if value_4183 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 4183的值！')
            if value_4233 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 Channel 4233的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need, u'Peak-to-Average Ratio – Result')


# 插入图片到word
def fccgsmwcdma_pictoword(casetype, casename, pic_name_list, expect_str):
    global_element.emitsingle.stateupdataSingle.emit('正在处理' + casetype + ' ' + casename + '图片……')
    # 遍历该项图片文件夹下所有图片
    dir_path = global_element.reporthandle_dict['datasource'].split('.')[0] + '/' + casetype + '/' + casename
    picture_name_list = global_element.ergodic_dir(dir_path)
    picture_count = len(picture_name_list)
    table_row_count = ((picture_count - 1) // 2 + 1) * 2
    table_column_count = 2
    document = Document(global_element.reporthandle_dict['reportaddr'])
    tab = document.add_table(rows=table_row_count, cols=table_column_count)

    # 处理GSM图片
    jushu = 1

    # 排序文件夹中图片名称
    picture_name_list_final = [x for x in pic_name_list if x in picture_name_list]

    for pic in picture_name_list_final:
        # 在表格中加入图片名称
        current_row = ((jushu - 1) // 2) * 2
        current_column = (jushu - 1) % 2
        cell = tab.cell(current_row, current_column)
        ph = cell.paragraphs[0]
        run = ph.add_run(pic)

        # 在表格中加入图片
        current_row = ((jushu - 1) // 2) * 2 + 1
        current_column = (jushu - 1) % 2
        cell = tab.cell(current_row, current_column)
        ph = cell.paragraphs[0]
        run = ph.add_run('')
        run.add_picture(dir_path + '/' + pic + '.BMP', width=Inches(3))

        jushu += 1

    expect_text = expect_str
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text.endswith(expect_text):
            target = paragraph
            break

    tbl, p = tab._tbl, target._p
    p.addnext(tbl)
    document.save(global_element.reporthandle_dict['reportaddr'])


# 处理gsm wcdma 26dB BW 和 99% BW数据到word
def bw_data_to_word():
    global_element.emitsingle.stateupdataSingle.emit('正在处理GSM & WCDMA OBW数据……')

    # 获取 GSM + WCDMA 的band总数，用以确定table的样式
    datasource = global_element.reporthandle_dict['datasource']
    datasource_wb = load_workbook(datasource)
    sheetname_list = datasource_wb.sheetnames
    gsm_band_list = []
    gprs_band_list = []
    egprs_band_list = []
    wcdma_band_list = []
    if 'GSM' in sheetname_list:
        sheet_gsm = datasource_wb['GSM']
        row_max_count = sheet_gsm.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_gsm.cell(row=row_count, column=8).value
            mode_value = sheet_gsm.cell(row=row_count, column=13).value
            item_value = sheet_gsm.cell(row=row_count, column=3).value
            if band_value not in gsm_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GSM' and \
                    item_value == '99% Occupied Bandwidth':
                gsm_band_list.append(band_value)
            elif band_value not in gprs_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GPRS' and \
                    item_value == '99% Occupied Bandwidth':
                gprs_band_list.append(band_value)
            elif band_value not in egprs_band_list and band_value in ['GSM850', 'PCS1900'] and \
                    mode_value == 'EGPRS' and item_value == '99% Occupied Bandwidth':
                egprs_band_list.append(band_value)

    if 'WCDMA' in sheetname_list:
        sheet_wcdma = datasource_wb['WCDMA']
        row_max_count = sheet_wcdma.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_wcdma.cell(row=row_count, column=8).value
            item_value = sheet_wcdma.cell(row=row_count, column=3).value
            if band_value not in wcdma_band_list and band_value in ['Band II', 'Band IV', 'Band V'] and \
                    item_value == '99% Occupied Bandwidth':
                wcdma_band_list.append(band_value)

    global_element.emitsingle.stateupdataSingle.emit('原始数据包含Band:' + '、'.join(gsm_band_list) + '、' +
                                                     '、'.join(wcdma_band_list))

    if len(egprs_band_list) > 0:
        # 提取出对应band数量的 EGPRS  power table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(egprs_band_list) + 9]

        # EGPRS power数据填入表内
        band_jishu = 0
        if 'GSM850' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'EGPRS850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128_26 = 'null'
            value_190_26 = 'null'
            value_251_26 = 'null'
            value_128_99 = 'null'
            value_190_99 = 'null'
            value_251_99 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '128' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_128_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '128' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_128_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '190' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_190_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '190' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_190_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '251' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_251_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '251' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_251_99 = sheet_gsm.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_need.cell(row + 1, coloumn).text = value_fre_190
            table_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_need.cell(row + 2, coloumn - 1).text = value_128_26
            table_need.cell(row + 2, coloumn).text = value_190_26
            table_need.cell(row + 2, coloumn + 1).text = value_251_26
            table_need.cell(row + 3, coloumn - 1).text = value_128_99
            table_need.cell(row + 3, coloumn).text = value_190_99
            table_need.cell(row + 3, coloumn + 1).text = value_251_99

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 EGPRS Channel 128的值！')
            if value_fre_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 EGPRS Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 EGPRS Channel 251的值！')

        if 'PCS1900' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'EGPRS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512_26 = 'null'
            value_661_26 = 'null'
            value_810_26 = 'null'
            value_512_99 = 'null'
            value_661_99 = 'null'
            value_810_99 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '512' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_512_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '512' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_512_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '661' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_661_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '661' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_661_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '810' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_810_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '810' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_810_99 = sheet_gsm.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_need.cell(row + 1, coloumn).text = value_fre_661
            table_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_need.cell(row + 2, coloumn - 1).text = value_512_26
            table_need.cell(row + 2, coloumn).text = value_661_26
            table_need.cell(row + 2, coloumn + 1).text = value_810_26
            table_need.cell(row + 3, coloumn - 1).text = value_512_99
            table_need.cell(row + 3, coloumn).text = value_661_99
            table_need.cell(row + 3, coloumn + 1).text = value_810_99

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 EGPRS Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 EGPRS Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 EGPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need,
                      u'26dB Bandwidth & Occupied Bandwidth – Result')

    if len(gprs_band_list) > 0:
        # 提取出对应band数量的 GPRS  BW table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(gprs_band_list) + 9]

        # GPRS BW数据填入表内
        band_jishu = 0
        if 'GSM850' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GPRS850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128_26 = 'null'
            value_190_26 = 'null'
            value_251_26 = 'null'
            value_128_99 = 'null'
            value_190_99 = 'null'
            value_251_99 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '128' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_128_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '128' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_128_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '190' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_190_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '190' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_190_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '251' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_251_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '251' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_251_99 = sheet_gsm.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_need.cell(row + 1, coloumn).text = value_fre_190
            table_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_need.cell(row + 2, coloumn - 1).text = value_128_26
            table_need.cell(row + 2, coloumn).text = value_190_26
            table_need.cell(row + 2, coloumn + 1).text = value_251_26
            table_need.cell(row + 3, coloumn - 1).text = value_128_99
            table_need.cell(row + 3, coloumn).text = value_190_99
            table_need.cell(row + 3, coloumn + 1).text = value_251_99

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GPRS Channel 128的值！')
            if value_fre_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GPRS Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GPRS Channel 251的值！')

        if 'PCS1900' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GPRS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512_26 = 'null'
            value_661_26 = 'null'
            value_810_26 = 'null'
            value_512_99 = 'null'
            value_661_99 = 'null'
            value_810_99 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '512' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_512_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '512' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_512_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '661' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_661_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '661' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_661_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '810' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_810_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '810' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_810_99 = sheet_gsm.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_need.cell(row + 1, coloumn).text = value_fre_661
            table_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_need.cell(row + 2, coloumn - 1).text = value_512_26
            table_need.cell(row + 2, coloumn).text = value_661_26
            table_need.cell(row + 2, coloumn + 1).text = value_810_26
            table_need.cell(row + 3, coloumn - 1).text = value_512_99
            table_need.cell(row + 3, coloumn).text = value_661_99
            table_need.cell(row + 3, coloumn + 1).text = value_810_99

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GPRS Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GPRS Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need,
                      u'26dB Bandwidth & Occupied Bandwidth – Result')

    # 将gsm wcdma bw数据导入word
    if len(gsm_band_list) + len(wcdma_band_list) > 0:
        # 提取出对应band数量的 gsm wcdma  BW table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(gsm_band_list) + len(wcdma_band_list) + 9]

        # gsm wcdma BW数据填入表内
        band_jishu = 0
        if 'GSM850' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GSM850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_128_26 = 'null'
            value_190_26 = 'null'
            value_251_26 = 'null'
            value_128_99 = 'null'
            value_190_99 = 'null'
            value_251_99 = 'null'
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '128' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_128_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '128' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_128_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '190' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_190_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '190' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_190_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '251' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_251_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '251' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_251_99 = sheet_gsm.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_128
            table_need.cell(row + 1, coloumn).text = value_fre_190
            table_need.cell(row + 1, coloumn + 1).text = value_fre_251
            table_need.cell(row + 2, coloumn - 1).text = value_128_26
            table_need.cell(row + 2, coloumn).text = value_190_26
            table_need.cell(row + 2, coloumn + 1).text = value_251_26
            table_need.cell(row + 3, coloumn - 1).text = value_128_99
            table_need.cell(row + 3, coloumn).text = value_190_99
            table_need.cell(row + 3, coloumn + 1).text = value_251_99

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GSM Channel 128的值！')
            if value_fre_190 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GSM Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GSM Channel 251的值！')

        if 'PCS1900' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'PCS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_512_26 = 'null'
            value_661_26 = 'null'
            value_810_26 = 'null'
            value_512_99 = 'null'
            value_661_99 = 'null'
            value_810_99 = 'null'
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '512' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_512_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '512' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_512_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '661' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_661_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '661' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_661_99 = sheet_gsm.cell(row=row_count, column=4).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '810' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_810_26 = sheet_gsm.cell(row=row_count, column=4).value
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '810' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_810_99 = sheet_gsm.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_512
            table_need.cell(row + 1, coloumn).text = value_fre_661
            table_need.cell(row + 1, coloumn + 1).text = value_fre_810
            table_need.cell(row + 2, coloumn - 1).text = value_512_26
            table_need.cell(row + 2, coloumn).text = value_661_26
            table_need.cell(row + 2, coloumn + 1).text = value_810_26
            table_need.cell(row + 3, coloumn - 1).text = value_512_99
            table_need.cell(row + 3, coloumn).text = value_661_99
            table_need.cell(row + 3, coloumn + 1).text = value_810_99

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GSM Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GSM Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 GSM Channel 810的值！')

        if 'Band II' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band II'
            table_need.cell(row, coloumn - 1).text = '9262'
            table_need.cell(row, coloumn).text = '9400'
            table_need.cell(row, coloumn + 1).text = '9538'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_9262_26 = 'null'
            value_9400_26 = 'null'
            value_9538_26 = 'null'
            value_9262_99 = 'null'
            value_9400_99 = 'null'
            value_9538_99 = 'null'
            value_fre_9262 = 'null'
            value_fre_9400 = 'null'
            value_fre_9538 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band II' and channelvalue == '9262' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_9262_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9262 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band II' and channelvalue == '9262' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_9262_99 = sheet_wcdma.cell(row=row_count, column=4).value
                elif bandvalue == 'Band II' and channelvalue == '9400' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_9400_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9400 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band II' and channelvalue == '9400' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_9400_99 = sheet_wcdma.cell(row=row_count, column=4).value
                elif bandvalue == 'Band II' and channelvalue == '9538' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_9538_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_9538 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band II' and channelvalue == '9538' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_9538_99 = sheet_wcdma.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_9262
            table_need.cell(row + 1, coloumn).text = value_fre_9400
            table_need.cell(row + 1, coloumn + 1).text = value_fre_9538
            table_need.cell(row + 2, coloumn - 1).text = value_9262_26
            table_need.cell(row + 2, coloumn).text = value_9400_26
            table_need.cell(row + 2, coloumn + 1).text = value_9538_26
            table_need.cell(row + 3, coloumn - 1).text = value_9262_99
            table_need.cell(row + 3, coloumn).text = value_9400_99
            table_need.cell(row + 3, coloumn + 1).text = value_9538_99

            table_need.cell(row + 2, coloumn - 2).text = '26dB OBW(MHz)'
            table_need.cell(row + 3, coloumn - 2).text = '99% OBW(MHz)'

            if value_fre_9262 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 9262的值！')
            if value_fre_9400 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 9400的值！')
            if value_fre_9538 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 9538的值！')

        if 'Band IV' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band IV'
            table_need.cell(row, coloumn - 1).text = '1312'
            table_need.cell(row, coloumn).text = '1413'
            table_need.cell(row, coloumn + 1).text = '1513'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_1312_26 = 'null'
            value_1413_26 = 'null'
            value_1513_26 = 'null'
            value_1312_99 = 'null'
            value_1413_99 = 'null'
            value_1513_99 = 'null'
            value_fre_1312 = 'null'
            value_fre_1413 = 'null'
            value_fre_1513 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band IV' and channelvalue == '1312' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_1312_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1312 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band IV' and channelvalue == '1312' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_1312_99 = sheet_wcdma.cell(row=row_count, column=4).value
                elif bandvalue == 'Band IV' and channelvalue == '1413' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_1413_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1413 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band IV' and channelvalue == '1413' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_1413_99 = sheet_wcdma.cell(row=row_count, column=4).value
                elif bandvalue == 'Band IV' and channelvalue == '1513' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_1513_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_1513 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band IV' and channelvalue == '1513' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_1513_99 = sheet_wcdma.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_1312
            table_need.cell(row + 1, coloumn).text = value_fre_1413
            table_need.cell(row + 1, coloumn + 1).text = value_fre_1513
            table_need.cell(row + 2, coloumn - 1).text = value_1312_26
            table_need.cell(row + 2, coloumn).text = value_1413_26
            table_need.cell(row + 2, coloumn + 1).text = value_1513_26
            table_need.cell(row + 3, coloumn - 1).text = value_1312_99
            table_need.cell(row + 3, coloumn).text = value_1413_99
            table_need.cell(row + 3, coloumn + 1).text = value_1513_99

            table_need.cell(row + 2, coloumn - 2).text = '26dB OBW(MHz)'
            table_need.cell(row + 3, coloumn - 2).text = '99% OBW(MHz)'

            if value_fre_1312 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 1312的值！')
            if value_fre_1413 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 1413的值！')
            if value_fre_1513 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 1513的值！')

        if 'Band V' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 5 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band V'
            table_need.cell(row, coloumn - 1).text = '4132'
            table_need.cell(row, coloumn).text = '4183'
            table_need.cell(row, coloumn + 1).text = '4233'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_4132_26 = 'null'
            value_4183_26 = 'null'
            value_4233_26 = 'null'
            value_4132_99 = 'null'
            value_4183_99 = 'null'
            value_4233_99 = 'null'
            value_fre_4132 = 'null'
            value_fre_4183 = 'null'
            value_fre_4233 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band V' and channelvalue == '4132' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_4132_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4132 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band V' and channelvalue == '4132' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_4132_99 = sheet_wcdma.cell(row=row_count, column=4).value
                elif bandvalue == 'Band V' and channelvalue == '4183' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_4183_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4183 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band V' and channelvalue == '4183' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_4183_99 = sheet_wcdma.cell(row=row_count, column=4).value
                elif bandvalue == 'Band V' and channelvalue == '4233' and \
                        itemvalue == '26dB Occupied Bandwidth':
                    value_4233_26 = sheet_wcdma.cell(row=row_count, column=4).value
                    value_fre_4233 = sheet_wcdma.cell(row=row_count, column=11).value
                elif bandvalue == 'Band V' and channelvalue == '4233' and \
                        itemvalue == '99% Occupied Bandwidth':
                    value_4233_99 = sheet_wcdma.cell(row=row_count, column=4).value

            table_need.cell(row + 1, coloumn - 1).text = value_fre_4132
            table_need.cell(row + 1, coloumn).text = value_fre_4183
            table_need.cell(row + 1, coloumn + 1).text = value_fre_4233
            table_need.cell(row + 2, coloumn - 1).text = value_4132_26
            table_need.cell(row + 2, coloumn).text = value_4183_26
            table_need.cell(row + 2, coloumn + 1).text = value_4233_26
            table_need.cell(row + 3, coloumn - 1).text = value_4132_99
            table_need.cell(row + 3, coloumn).text = value_4183_99
            table_need.cell(row + 3, coloumn + 1).text = value_4233_99

            table_need.cell(row + 2, coloumn - 2).text = '26dB OBW(MHz)'
            table_need.cell(row + 3, coloumn - 2).text = '99% OBW(MHz)'

            if value_fre_4132 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 4132的值！')
            if value_fre_4183 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 4183的值！')
            if value_fre_4233 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：       没有获取有效的 WCDMA Channel 4233的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need,
                      u'26dB Bandwidth & Occupied Bandwidth – Result')


# 处理gsm & wcdma Band edge数据到word报告
def fcc_gsmwcdma_bandedge_to_word():
    global_element.emitsingle.stateupdataSingle.emit('正在处理GSM & WCDMA Band edge数据……')

    # 获取 GSM + WCDMA 的band总数，用以确定table的样式
    datasource = global_element.reporthandle_dict['datasource']
    datasource_wb = load_workbook(datasource)
    sheetname_list = datasource_wb.sheetnames
    gsm_band_list = []
    gprs_band_list = []
    egprs_band_list = []
    wcdma_band_list = []
    if 'GSM' in sheetname_list:
        sheet_gsm = datasource_wb['GSM']
        row_max_count = sheet_gsm.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_gsm.cell(row=row_count, column=8).value
            mode_value = sheet_gsm.cell(row=row_count, column=13).value
            item_value = sheet_gsm.cell(row=row_count, column=3).value
            if band_value not in gsm_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GSM' and \
                    item_value == 'Band edge':
                gsm_band_list.append(band_value)
            elif band_value not in gprs_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GPRS' and \
                    item_value == 'Band edge':
                gprs_band_list.append(band_value)
            elif band_value not in egprs_band_list and band_value in ['GSM850', 'PCS1900'] and \
                    mode_value == 'EGPRS' and item_value == 'Band edge':
                egprs_band_list.append(band_value)

    if 'WCDMA' in sheetname_list:
        sheet_wcdma = datasource_wb['WCDMA']
        row_max_count = sheet_wcdma.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_wcdma.cell(row=row_count, column=8).value
            item_value = sheet_wcdma.cell(row=row_count, column=3).value
            if band_value not in wcdma_band_list and band_value in ['Band II', 'Band IV', 'Band V'] and \
                    item_value == 'Band edge':
                wcdma_band_list.append(band_value)

    global_element.emitsingle.stateupdataSingle.emit('原始数据包含Band:' + '、'.join(gsm_band_list) + '、' +
                                                     '、'.join(wcdma_band_list))

    if len(egprs_band_list) > 0:
        # 提取出对应band数量的 EGPRS  Band edge table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(egprs_band_list) + 14]

        # EGPRS 数据填入表内
        band_jishu = 0
        if 'GSM850' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'EGPRS850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_128 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '128' and \
                        itemvalue == 'Band edge':
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '251' and \
                        itemvalue == 'Band edge':
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_251]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 1, coloumn).text = value_judgement_251

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 128的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 251的值！')

        if 'PCS1900' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'EGPRS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_512 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '512' and \
                        itemvalue == 'Band edge':
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '810' and \
                        itemvalue == 'Band edge':
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_810]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 1, coloumn).text = value_judgement_810

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 512的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need, u'Conducted Band Edge – Result')

    if len(gprs_band_list) > 0:
        # 提取出对应band数量的 GPRS  Band edge table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(gprs_band_list) + 14]

        # GPRS 数据填入表内
        band_jishu = 0
        if 'GSM850' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GPRS850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_128 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '128' and \
                        itemvalue == 'Band edge':
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '251' and \
                        itemvalue == 'Band edge':
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_251]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 1, coloumn).text = value_judgement_251

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 128的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 251的值！')

        if 'PCS1900' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GPRS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_512 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '512' and \
                        itemvalue == 'Band edge':
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '810' and \
                        itemvalue == 'Band edge':
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_810]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 1, coloumn).text = value_judgement_810

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 512的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need, u'Conducted Band Edge – Result')

    if (len(gsm_band_list) + len(wcdma_band_list)) > 0:
        # 提取出对应band数量的 GSM WCDMA  Band edge table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(gsm_band_list) + len(wcdma_band_list) + 14]

        # GSM WCDMA 数据填入表内
        band_jishu = 0
        if 'GSM850' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GSM850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_128 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '128' and \
                        itemvalue == 'Band edge':
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '251' and \
                        itemvalue == 'Band edge':
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_251]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 1, coloumn).text = value_judgement_251

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 128的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 251的值！')

        if 'PCS1900' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'PCS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_512 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '512' and \
                        itemvalue == 'Band edge':
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '810' and \
                        itemvalue == 'Band edge':
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_810]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 1, coloumn).text = value_judgement_810

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 512的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 810的值！')

        if 'Band II' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band II'
            table_need.cell(row, coloumn - 1).text = '9262'
            table_need.cell(row, coloumn).text = '9538'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_fre_9262 = 'null'
            value_fre_9538 = 'null'
            value_judgement_9262 = 'null'
            value_judgement_9538 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band II' and channelvalue == '9262' and \
                        itemvalue == 'Band edge':
                    value_fre_9262 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_9262 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band II' and channelvalue == '9538' and \
                        itemvalue == 'Band edge':
                    value_fre_9538 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_9538 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_9262, value_judgement_9538]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_9262
            table_need.cell(row + 1, coloumn).text = value_judgement_9538

            if value_fre_9262 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 9262的值！')
            if value_fre_9538 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 9538的值！')

        if 'Band IV' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band IV'
            table_need.cell(row, coloumn - 1).text = '1312'
            table_need.cell(row, coloumn).text = '1513'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_fre_1312 = 'null'
            value_fre_1513 = 'null'
            value_judgement_1312 = 'null'
            value_judgement_1513 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band IV' and channelvalue == '1312' and \
                        itemvalue == 'Band edge':
                    value_fre_1312 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_1312 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band IV' and channelvalue == '1513' and \
                        itemvalue == 'Band edge':
                    value_fre_1513 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_1513 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_1312, value_judgement_1513]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_1312
            table_need.cell(row + 1, coloumn).text = value_judgement_1513

            if value_fre_1312 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 1312的值！')
            if value_fre_1513 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 1513的值！')

        if 'Band V' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 2 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band V'
            table_need.cell(row, coloumn - 1).text = '4132'
            table_need.cell(row, coloumn).text = '4233'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_fre_4132 = 'null'
            value_fre_4233 = 'null'
            value_judgement_4132 = 'null'
            value_judgement_4233 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band V' and channelvalue == '4132' and \
                        itemvalue == 'Band edge':
                    value_fre_4132 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_4132 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band V' and channelvalue == '4233' and \
                        itemvalue == 'Band edge':
                    value_fre_4233 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_4233 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_4132, value_judgement_4233]:
                global_element.fcc_gsm_wcdma_bandedge_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_4132
            table_need.cell(row + 1, coloumn).text = value_judgement_4233

            if value_fre_4132 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 4132的值！')
            if value_fre_4233 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 4233的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need, u'Conducted Band Edge – Result')


# 处理gsm & wcdma Band edge数据到word报告
def fcc_gsmwcdma_cse_to_word():
    global_element.emitsingle.stateupdataSingle.emit('正在处理GSM & WCDMA Conducted Spurious Emissions数据……')

    # 获取 GSM + WCDMA 的band总数，用以确定table的样式
    datasource = global_element.reporthandle_dict['datasource']
    datasource_wb = load_workbook(datasource)
    sheetname_list = datasource_wb.sheetnames
    gsm_band_list = []
    gprs_band_list = []
    egprs_band_list = []
    wcdma_band_list = []
    if 'GSM' in sheetname_list:
        sheet_gsm = datasource_wb['GSM']
        row_max_count = sheet_gsm.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_gsm.cell(row=row_count, column=8).value
            mode_value = sheet_gsm.cell(row=row_count, column=13).value
            item_value = sheet_gsm.cell(row=row_count, column=3).value
            if band_value not in gsm_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GSM' and \
                    item_value == 'Conducted Spurious emissions':
                gsm_band_list.append(band_value)
            elif band_value not in gprs_band_list and band_value in ['GSM850', 'PCS1900'] and mode_value == 'GPRS' and \
                    item_value == 'Conducted Spurious emissions':
                gprs_band_list.append(band_value)
            elif band_value not in egprs_band_list and band_value in ['GSM850', 'PCS1900'] and \
                    mode_value == 'EGPRS' and item_value == 'Conducted Spurious emissions':
                egprs_band_list.append(band_value)

    if 'WCDMA' in sheetname_list:
        sheet_wcdma = datasource_wb['WCDMA']
        row_max_count = sheet_wcdma.max_row
        for row_count in range(2, row_max_count + 1):
            band_value = sheet_wcdma.cell(row=row_count, column=8).value
            item_value = sheet_wcdma.cell(row=row_count, column=3).value
            if band_value not in wcdma_band_list and band_value in ['Band II', 'Band IV', 'Band V'] and \
                    item_value == 'Conducted Spurious emissions':
                wcdma_band_list.append(band_value)

    global_element.emitsingle.stateupdataSingle.emit('原始数据包含Band:' + '、'.join(gsm_band_list) + '、' +
                                                     '、'.join(wcdma_band_list))

    if len(egprs_band_list) > 0:
        # 提取出对应band数量的 EGPRS  Conducted Spurious emissions table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(egprs_band_list) + 19]

        # EGPRS 数据填入表内
        band_jishu = 0
        if 'GSM850' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'EGPRS850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '128' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '190' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_190 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'EGPRS' and channelvalue == '251' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_190, value_judgement_251]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 1, coloumn).text = value_judgement_190
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_251

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 128的值！')
                if value_fre_190 == 'null':
                    global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 251的值！')

        if 'PCS1900' in egprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'EGPRS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '512' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '661' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_661 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'EGPRS' and channelvalue == '810' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_661, value_judgement_810]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 1, coloumn).text = value_judgement_661
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_810

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 EGPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need,
                      u'Conducted Spurious Emission – Result')

    if len(gprs_band_list) > 0:
        # 提取出对应band数量的 GPRS  Conducted Spurious emissions table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(egprs_band_list) + 19]

        # GPRS 数据填入表内
        band_jishu = 0
        if 'GSM850' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GPRS850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '128' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '190' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_190 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GPRS' and channelvalue == '251' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_190, value_judgement_251]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 1, coloumn).text = value_judgement_190
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_251

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 128的值！')
                if value_fre_190 == 'null':
                    global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 251的值！')

        if 'PCS1900' in gprs_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GPRS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '512' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '661' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_661 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GPRS' and channelvalue == '810' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_661, value_judgement_810]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 1, coloumn).text = value_judgement_661
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_810

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GPRS Channel 810的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need,
                      u'Conducted Spurious Emission – Result')

    if len(gsm_band_list) + len(wcdma_band_list) > 0:
        # 提取出对应band数量的 GSM WCDMA Conducted Spurious emissions table模板
        document = Document('Config/Report mask/gsm & wcdma table mudo.docx')
        tables = document.tables
        table_need = tables[len(gsm_band_list) + len(wcdma_band_list) + 19]

        # GSM WCDMA 数据填入表内
        band_jishu = 0
        if 'GSM850' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'GSM850'
            table_need.cell(row, coloumn - 1).text = '128'
            table_need.cell(row, coloumn).text = '190'
            table_need.cell(row, coloumn + 1).text = '251'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_128 = 'null'
            value_fre_190 = 'null'
            value_fre_251 = 'null'
            value_judgement_128 = 'null'
            value_judgement_190 = 'null'
            value_judgement_251 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '128' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_128 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_128 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '190' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_190 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_190 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'GSM850' and modevalue == 'GSM' and channelvalue == '251' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_251 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_251 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_128, value_judgement_190, value_judgement_251]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_128
            table_need.cell(row + 1, coloumn).text = value_judgement_190
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_251

            if value_fre_128 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 128的值！')
                if value_fre_190 == 'null':
                    global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 190的值！')
            if value_fre_251 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 251的值！')

        if 'PCS1900' in gsm_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'PCS1900'
            table_need.cell(row, coloumn - 1).text = '512'
            table_need.cell(row, coloumn).text = '661'
            table_need.cell(row, coloumn + 1).text = '810'

            sheet_gsm = datasource_wb['GSM']
            row_max_count = sheet_gsm.max_row
            value_fre_512 = 'null'
            value_fre_661 = 'null'
            value_fre_810 = 'null'
            value_judgement_512 = 'null'
            value_judgement_661 = 'null'
            value_judgement_810 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_gsm.cell(row=row_count, column=8).value
                modevalue = sheet_gsm.cell(row=row_count, column=13).value
                channelvalue = sheet_gsm.cell(row=row_count, column=9).value
                itemvalue = sheet_gsm.cell(row=row_count, column=3).value
                if bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '512' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_512 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_512 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '661' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_661 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_661 = sheet_gsm.cell(row=row_count, column=7).value
                elif bandvalue == 'PCS1900' and modevalue == 'GSM' and channelvalue == '810' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_810 = sheet_gsm.cell(row=row_count, column=11).value
                    value_judgement_810 = sheet_gsm.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_512, value_judgement_661, value_judgement_810]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_512
            table_need.cell(row + 1, coloumn).text = value_judgement_661
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_810

            if value_fre_512 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 512的值！')
            if value_fre_661 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 661的值！')
            if value_fre_810 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 GSM Channel 810的值！')

        if 'Band II' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band II'
            table_need.cell(row, coloumn - 1).text = '9262'
            table_need.cell(row, coloumn).text = '9400'
            table_need.cell(row, coloumn + 1).text = '9538'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_fre_9262 = 'null'
            value_fre_9400 = 'null'
            value_fre_9538 = 'null'
            value_judgement_9262 = 'null'
            value_judgement_9400 = 'null'
            value_judgement_9538 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band II' and channelvalue == '9262' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_9262 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_9262 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band II' and channelvalue == '9400' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_9400 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_9400 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band II' and channelvalue == '9538' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_9538 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_9538 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_9262, value_judgement_9400, value_judgement_9538]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_9262
            table_need.cell(row + 1, coloumn).text = value_judgement_9400
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_9538

            if value_fre_9262 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 9262的值！')
            if value_fre_9400 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 9400的值！')
            if value_fre_9538 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 9538的值！')

        if 'Band IV' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band IV'
            table_need.cell(row, coloumn - 1).text = '1312'
            table_need.cell(row, coloumn).text = '1413'
            table_need.cell(row, coloumn + 1).text = '1513'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_fre_1312 = 'null'
            value_fre_1413 = 'null'
            value_fre_1513 = 'null'
            value_judgement_1312 = 'null'
            value_judgement_1413 = 'null'
            value_judgement_1513 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band IV' and channelvalue == '1312' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_1312 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_1312 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band IV' and channelvalue == '1413' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_1413 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_1413 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band IV' and channelvalue == '1513' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_1513 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_1513 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_1312, value_judgement_1413, value_judgement_1513]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_1312
            table_need.cell(row + 1, coloumn).text = value_judgement_1413
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_1513

            if value_fre_1312 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 1312的值！')
            if value_fre_1413 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 1413的值！')
            if value_fre_1513 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 1513的值！')

        if 'Band V' in wcdma_band_list:
            band_jishu += 1
            row = ((band_jishu - 1) // 2) * 3 + 1
            coloumn = ((band_jishu - 1) % 2) * 3 + 2
            table_need.cell(row - 1, coloumn - 1).text = 'Band V'
            table_need.cell(row, coloumn - 1).text = '4132'
            table_need.cell(row, coloumn).text = '4183'
            table_need.cell(row, coloumn + 1).text = '4233'

            sheet_wcdma = datasource_wb['WCDMA']
            row_max_count = sheet_wcdma.max_row
            value_fre_4132 = 'null'
            value_fre_4183 = 'null'
            value_fre_4233 = 'null'
            value_judgement_4132 = 'null'
            value_judgement_4183 = 'null'
            value_judgement_4233 = 'null'
            for row_count in range(2, row_max_count + 1):
                bandvalue = sheet_wcdma.cell(row=row_count, column=8).value
                channelvalue = sheet_wcdma.cell(row=row_count, column=9).value
                itemvalue = sheet_wcdma.cell(row=row_count, column=3).value
                if bandvalue == 'Band V' and channelvalue == '4132' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_4132 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_4132 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band V' and channelvalue == '4183' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_4183 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_4183 = sheet_wcdma.cell(row=row_count, column=7).value
                elif bandvalue == 'Band V' and channelvalue == '4233' and \
                        itemvalue == 'Conducted Spurious emissions':
                    value_fre_4233 = sheet_wcdma.cell(row=row_count, column=11).value
                    value_judgement_4233 = sheet_wcdma.cell(row=row_count, column=7).value

            if 'Failed' in [value_judgement_4132, value_judgement_4183, value_judgement_4233]:
                global_element.fcc_gsm_wcdma_cse_judgement = 'Failed'

            table_need.cell(row + 1, coloumn - 1).text = value_judgement_4132
            table_need.cell(row + 1, coloumn).text = value_judgement_4183
            table_need.cell(row + 1, coloumn + 1).text = value_judgement_4233

            if value_fre_4132 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 4132的值！')
            if value_fre_4183 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 4183的值！')
            if value_fre_4233 == 'null':
                global_element.emitsingle.stateupdataSingle.emit('错误：  没有获取有效的 WCDMA Channel 4233的值！')

        tabletoreport(global_element.reporthandle_dict['reportaddr'], table_need,
                      u'Conducted Spurious Emission – Result')


# # 插入图片到word
# def pictoword_ww(dirpath, docpath):
#
#     dir_path = dirpath
#     picture_name_list = global_element.ergodic_dir(dir_path)
#     picture_count = len(picture_name_list)
#     table_row_count = ((picture_count - 1) // 2 + 1) * 2
#     table_column_count = 2
#     document = Document(docpath)
#     tab = document.add_table(rows=table_row_count, cols=table_column_count)
#
#     # 处理GSM图片
#     jushu = 1
#
#     picture_name_list.sort()
#     picture_name_list = ['FDD-LTE 2 BW1.4MHz Channel18607 RB 1 16-QAM', 'FDD-LTE 2 BW1.4MHz Channel18607 RB 1 QPSK', 'FDD-LTE 2 BW1.4MHz Channel18900 RB 1 16-QAM', 'FDD-LTE 2 BW1.4MHz Channel18900 RB 1 QPSK', 'FDD-LTE 2 BW1.4MHz Channel19193 RB 1 16-QAM', 'FDD-LTE 2 BW1.4MHz Channel19193 RB 1 QPSK', 'FDD-LTE 2 BW3MHz Channel18615 RB 1 16-QAM', 'FDD-LTE 2 BW3MHz Channel18615 RB 1 QPSK', 'FDD-LTE 2 BW3MHz Channel18900 RB 1 16-QAM', 'FDD-LTE 2 BW3MHz Channel18900 RB 1 QPSK', 'FDD-LTE 2 BW3MHz Channel19185 RB 1 16-QAM', 'FDD-LTE 2 BW3MHz Channel19185 RB 1 QPSK', 'FDD-LTE 2 BW5MHz Channel18625 RB 1 16-QAM', 'FDD-LTE 2 BW5MHz Channel18625 RB 1 QPSK', 'FDD-LTE 2 BW5MHz Channel18900 RB 1 16-QAM', 'FDD-LTE 2 BW5MHz Channel18900 RB 1 QPSK', 'FDD-LTE 2 BW5MHz Channel19175 RB 1 16-QAM', 'FDD-LTE 2 BW5MHz Channel19175 RB 1 QPSK', 'FDD-LTE 2 BW10MHz Channel18650 RB 1 16-QAM', 'FDD-LTE 2 BW10MHz Channel18650 RB 1 QPSK', 'FDD-LTE 2 BW10MHz Channel18900 RB 1 16-QAM', 'FDD-LTE 2 BW10MHz Channel18900 RB 1 QPSK', 'FDD-LTE 2 BW10MHz Channel19150 RB 1 16-QAM', 'FDD-LTE 2 BW10MHz Channel19150 RB 1 QPSK', 'FDD-LTE 2 BW15MHz Channel18675 RB 1 16-QAM', 'FDD-LTE 2 BW15MHz Channel18675 RB 1 QPSK', 'FDD-LTE 2 BW15MHz Channel18900 RB 1 16-QAM', 'FDD-LTE 2 BW15MHz Channel18900 RB 1 QPSK', 'FDD-LTE 2 BW15MHz Channel19125 RB 1 16-QAM', 'FDD-LTE 2 BW15MHz Channel19125 RB 1 QPSK', 'FDD-LTE 2 BW20MHz Channel18700 RB 1 16-QAM', 'FDD-LTE 2 BW20MHz Channel18700 RB 1 QPSK', 'FDD-LTE 2 BW20MHz Channel18900 RB 1 16-QAM', 'FDD-LTE 2 BW20MHz Channel18900 RB 1 QPSK', 'FDD-LTE 2 BW20MHz Channel19100 RB 1 16-QAM', 'FDD-LTE 2 BW20MHz Channel19100 RB 1 QPSK', 'FDD-LTE 4 BW1.4MHz Channel19957 RB 1 16-QAM', 'FDD-LTE 4 BW1.4MHz Channel19957 RB 1 QPSK', 'FDD-LTE 4 BW1.4MHz Channel20175 RB 1 16-QAM', 'FDD-LTE 4 BW1.4MHz Channel20175 RB 1 QPSK', 'FDD-LTE 4 BW1.4MHz Channel20393 RB 1 16-QAM', 'FDD-LTE 4 BW1.4MHz Channel20393 RB 1 QPSK', 'FDD-LTE 4 BW3MHz Channel19965 RB 1 16-QAM', 'FDD-LTE 4 BW3MHz Channel19965 RB 1 QPSK', 'FDD-LTE 4 BW3MHz Channel20175 RB 1 16-QAM', 'FDD-LTE 4 BW3MHz Channel20175 RB 1 QPSK', 'FDD-LTE 4 BW3MHz Channel20385 RB 1 16-QAM', 'FDD-LTE 4 BW3MHz Channel20385 RB 1 QPSK', 'FDD-LTE 4 BW5MHz Channel19975 RB 1 16-QAM', 'FDD-LTE 4 BW5MHz Channel19975 RB 1 QPSK', 'FDD-LTE 4 BW5MHz Channel20175 RB 1 16-QAM', 'FDD-LTE 4 BW5MHz Channel20175 RB 1 QPSK', 'FDD-LTE 4 BW5MHz Channel20375 RB 1 16-QAM', 'FDD-LTE 4 BW5MHz Channel20375 RB 1 QPSK', 'FDD-LTE 4 BW10MHz Channel20000 RB 1 16-QAM', 'FDD-LTE 4 BW10MHz Channel20000 RB 1 QPSK', 'FDD-LTE 4 BW10MHz Channel20175 RB 1 16-QAM', 'FDD-LTE 4 BW10MHz Channel20175 RB 1 QPSK', 'FDD-LTE 4 BW10MHz Channel20350 RB 1 16-QAM', 'FDD-LTE 4 BW10MHz Channel20350 RB 1 QPSK', 'FDD-LTE 4 BW15MHz Channel20025 RB 1 16-QAM', 'FDD-LTE 4 BW15MHz Channel20025 RB 1 QPSK', 'FDD-LTE 4 BW15MHz Channel20175 RB 1 16-QAM', 'FDD-LTE 4 BW15MHz Channel20175 RB 1 QPSK', 'FDD-LTE 4 BW15MHz Channel20325 RB 1 16-QAM', 'FDD-LTE 4 BW15MHz Channel20325 RB 1 QPSK', 'FDD-LTE 4 BW20MHz Channel20050 RB 1 16-QAM', 'FDD-LTE 4 BW20MHz Channel20050 RB 1 QPSK', 'FDD-LTE 4 BW20MHz Channel20175 RB 1 16-QAM', 'FDD-LTE 4 BW20MHz Channel20175 RB 1 QPSK', 'FDD-LTE 4 BW20MHz Channel20300 RB 1 16-QAM', 'FDD-LTE 4 BW20MHz Channel20300 RB 1 QPSK']
#     # print(picture_name_list)
#
#     for pic in picture_name_list:
#         # 在表格中加入图片名称
#         current_row = ((jushu - 1) // 2) * 2
#         current_column = (jushu - 1) % 2
#         cell = tab.cell(current_row, current_column)
#         ph = cell.paragraphs[0]
#         run = ph.add_run(pic)
#
#         # 在表格中加入图片
#         current_row = ((jushu - 1) // 2) * 2 + 1
#         current_column = (jushu - 1) % 2
#         cell = tab.cell(current_row, current_column)
#         ph = cell.paragraphs[0]
#         run = ph.add_run('')
#         run.add_picture(dir_path + '/' + pic + '.BMP', width=Inches(3))
#
#         jushu += 1
#
#
#     document.save(docpath)